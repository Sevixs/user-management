#!/usr/bin/env python3
"""
文件上传漏洞复现及修复报告 — Word 文档生成脚本
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="微软雅黑", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size: run.font.size = Pt(size)
    run.font.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        set_run_font(r, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return h


def add_para(doc, text, bold=False, size=11, color=None, space_after=6, first_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(20)
    if first_indent: p.paragraph_format.first_line_indent = Cm(first_indent)
    return p


def add_code(doc, code, indent=1.0):
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)


def add_separator(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)


def make_table(doc, headers, rows, col_widths=None, header_color="C0392B"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]; cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
            if ri % 2 == 1: set_cell_shading(cell, "F2F2F2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph()


def generate_report(output_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    for m in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, m, Cm(2.5))

    # ═══ 封面 ═══
    for _ in range(6): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Flask 用户信息管理平台"); set_run_font(run, size=26, bold=True, color=(192, 57, 43))
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run("文件上传漏洞复现及修复报告"); set_run_font(run, size=20, color=(89, 89, 89))
    doc.add_paragraph()
    line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━" * 40); set_run_font(run, size=12, color=(200, 200, 200))
    doc.add_paragraph()
    for t in ["实验性质：文件上传漏洞复现与安全加固实训", "目标系统：Flask 用户信息管理系统（头像上传）",
              "文档版本：V1.0 终审版", "生成日期：2026 年 7 月"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(t); set_run_font(run, size=12, color=(100, 100, 100))
    doc.add_page_break()

    # ═══ 目录 ═══
    add_heading_styled(doc, "目  录", level=1); doc.add_paragraph()
    for t in ["一、实验概述", "二、漏洞总览", "三、漏洞 1：路径遍历文件覆盖漏洞",
              "四、漏洞 2：无后缀白名单脚本后门上传", "五、漏洞 3：MIME 伪装绕过校验",
              "六、漏洞 4：原始文件名存储文件覆盖", "七、漏洞 5：存储型 XSS 漏洞",
              "八、漏洞 6：单文件大小限制缺失 DoS", "九、漏洞 7：上传目录匿名访问",
              "十、修复方案汇总与对比", "十一、修复后安全验证", "十二、总结与心得"]:
        p = doc.add_paragraph(); p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(3), Pt(21)
        run = p.add_run(t); set_run_font(run, size=11)
    doc.add_page_break()

    # ═══ 一、实验概述 ═══
    add_heading_styled(doc, "一、实验概述", level=1); add_separator(doc)
    add_para(doc, "本次实验针对 Flask 用户信息管理系统的头像上传功能中存在的 7 项高危文件上传漏洞进行手工复现与安全修复。该功能初始版本未做任何安全校验，包括不限制文件后缀、不校验文件内容、使用原始文件名存储、无访问鉴权等，攻击者可利用这些缺陷实现路径遍历、脚本后门上传、文件覆盖、XSS 攻击、磁盘耗尽拒绝服务等危害。", first_indent=0.74)
    add_para(doc, "涉及漏洞类型：", bold=True)
    add_para(doc, "（1）路径遍历与文件覆盖；（2）无后缀白名单 — 脚本后门上传；（3）MIME 伪装绕过；（4）原始文件名存储覆盖；（5）存储型 XSS；（6）单文件大小限制缺失 — DoS；（7）上传目录匿名公开访问", first_indent=0.74)

    # ═══ 二、漏洞总览 ═══
    add_heading_styled(doc, "二、漏洞总览", level=1); add_separator(doc)
    make_table(doc, ["编号", "漏洞名称", "类型", "危害等级"],
               [["VUL-UPLOAD-01", "路径遍历文件覆盖漏洞", "路径穿越/覆盖", "高危"],
                ["VUL-UPLOAD-02", "无后缀白名单脚本后门上传", "任意文件上传", "高危"],
                ["VUL-UPLOAD-03", "MIME 伪装绕过校验", "类型校验绕过", "高危"],
                ["VUL-UPLOAD-04", "原始文件名存储文件覆盖", "文件覆盖", "中危"],
                ["VUL-UPLOAD-05", "文件名无转义存储型 XSS", "XSS", "中危"],
                ["VUL-UPLOAD-06", "单文件大小缺失 DoS", "拒绝服务", "中危"],
                ["VUL-UPLOAD-07", "上传目录匿名公开访问", "未授权访问", "高危"]],
               col_widths=[3, 5, 3.5, 2])
    add_para(doc, "漏洞根因：未校验文件名、未校验文件后缀、未校验文件内容、使用原始文件名存储、无文件大小限制、无访问鉴权，用户输入完全可控。", bold=True)

    # ═══ 漏洞 1-7 ═══
    vulns = [
        ("三", "路径遍历文件覆盖漏洞", "搜索接口",
         "上传功能直接使用用户输入的原始文件名拼接路径保存，未做路径穿越字符过滤。攻击者可在 filename 参数中插入 ../ 符号，使文件保存到上层目录，覆盖 app.py 等关键源码文件。",
         "修改数据包内文件 filename 参数值为 ../app.py，其余参数不变。",
         "页面提示上传成功，返回静态资源 URL；服务器 static/uploads 上层目录 app.py 被上传文件覆盖，后端业务逻辑被篡改。",
         "使用 werkzeug.utils.secure_filename 清洗文件路径，过滤 ../ / 等路径穿越字符；完全弃用原始文件名，采用 UUID 随机字符串 + 合法图片后缀重命名存储。",
         'from werkzeug.utils import secure_filename\nimport uuid\n\nsafe_name = secure_filename(original_filename)\n# 弃用原始文件名，UUID 重命名\nsafe_filename = f"{uuid.uuid4().hex}.{image_type}"\nsave_path = os.path.join(UPLOAD_DIR, safe_filename)\nfile.save(save_path)',
         "secure_filename 移除../等路径字符；UUID 生成不可预测的唯一文件名，彻底杜绝目录遍历和文件覆盖。"),

        ("四", "无后缀白名单脚本后门上传", "上传接口",
         "上传接口未限制文件后缀类型，任意后缀（.php、.py、.js、.html 等）均可上传。攻击者可上传 WebShell 后门文件，通过浏览器直接访问执行恶意代码。",
         "本地新建 shell.php 一句话后门文件。选择该 php 文件上传，Burp 拦截确认 filename 为 shell.php，直接发送。",
         "文件上传成功存入 static/uploads/shell.php；匿名访问 URL 可正常执行 PHP 后门代码，实现服务器远程控制。",
         "配置图片后缀白名单，仅允许 jpg、jpeg、png、gif 四种格式；黑名单拦截 php、py、js、sh、sql、html 等所有可执行脚本后缀。",
         'ALLOWED_EXTENSIONS = {"jpg", "jpeg", "png", "gif"}\n\ndef allowed_file(filename):\n    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""\n    return ext in ALLOWED_EXTENSIONS\n\nif not allowed_file(original_filename):\n    error = "仅允许上传 jpg/jpeg/png/gif 格式的图片"',
         "白名单机制只放行已知安全的图片格式，任意可执行脚本后缀被直接拒绝，未经白名单允许的文件无法保存。"),

        ("五", "MIME 伪装绕过校验", "上传接口",
         "后端仅依赖前端请求头 Content-Type 判断文件类型，攻击者可篡改 Content-Type 为 image/png，将 PHP 后门伪装成图片上传。",
         "准备 shell.php 恶意脚本，Burp 拦截修改两处：filename=shell.php、Content-Type 改为 image/png。",
         "后端仅校验请求头 MIME，未校验文件真实二进制，脚本上传成功；匿名访问 URL 可执行后门代码。",
         "读取文件二进制头部校验真实 MIME 类型（魔数校验），不依赖前端 Content-Type 请求头。仅放行纯图片二进制文件。",
         'def check_image_magic(file_stream):\n    magic = file_stream.read(8)\n    file_stream.seek(0)\n    if magic[:3] == b"\\xff\\xd8\\xff": return True, "jpeg"\n    elif magic[:4] == b"\\x89PNG": return True, "png"\n    elif magic[:3] == b"GIF": return True, "gif"\n    else: return False, None',
         "魔数校验从文件二进制内容判别真实类型，攻击者篡改请求头无效；图片马内容不含图片魔数也被拦截。"),

        ("六", "原始文件名存储文件覆盖", "上传接口",
         "直接使用用户上传的原始文件名保存，不同用户上传同名文件时后上传的会覆盖先上传的。攻击者可制作恶意文件上传覆盖系统已有资源。",
         "提前确认网站静态目录存在 logo.png，制作同名恶意 png 文件上传修改文件名为 logo.png。",
         "恶意文件直接覆盖原有 logo.png，网站页面加载篡改后的图片。",
         "完全弃用用户原始文件名，采用 UUID 随机字符串 + 合法图片后缀生成唯一存储文件名，杜绝同名文件覆盖。",
         'safe_filename = f"{uuid.uuid4().hex}.{image_type}"',
         "UUID 每次生成全局唯一随机字符串，不同用户/不同次上传的文件名永不重复，从根源上杜绝文件覆盖。"),

        ("七", "文件名无转义存储型 XSS", "上传接口/展示页",
         "上传页面直接渲染用户提供的原始文件名，未做 HTML 转义处理。攻击者可将文件名构造为包含 JavaScript 脚本的内容，触发 XSS。",
         "Burp 拦截上传请求，修改文件名为 <script>alert(document.cookie)</script>.png。",
         "页面直接渲染原始恶意文件名，未转义 HTML 特殊字符，打开页面自动弹出 Cookie 弹窗。",
         "前端页面通过 Jinja2 自动转义引擎（autoescape）输出文件名，同时对文件名中的 HTML 特殊字符进行后端过滤。",
         '# 后端过滤函数\ndef sanitize_display_text(text):\n    replacements = {"&": "&amp;", "<": "&lt;", ">": "&gt;",\n                    "\'": "&#x27;", \'"\': "&quot;"}\n    for old, new in replacements.items():\n        text = text.replace(old, new)\n    return text\n\n# Jinja2 模板中 {{ display_filename }} 自动转义',
         "后端过滤 + Jinja2 autoescape 双重防护，HTML 特殊字符被转为实体编码，浏览器不会解析执行脚本。"),

        ("八", "单文件大小限制缺失 DoS", "上传接口",
         "仅设置全局 MAX_CONTENT_LENGTH=16MB，未对单文件大小做限制。攻击者可批量上传超大文件耗尽服务器磁盘空间。",
         "制作 15MB 空白超大文件，登录上传页面抓包上传并批量重复发送。",
         "后端仅限制全局 16MB，未限制单文件大小，超大文件全部上传成功；多次上传后服务器磁盘空间被大量占用，网站加载卡顿。",
         "新增单文件上限 5MB 校验，读取文件字节大小提前拦截超大恶意文件。",
         'MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB\n\nfile.seek(0, os.SEEK_END)\nfile_size = file.tell()\nfile.seek(0)\nif file_size > MAX_FILE_SIZE:\n    error = f"文件大小超过限制（最大 5MB）"',
         "在文件保存前读取字节数进行大小校验，超过 5MB 直接拒绝，防止超量文件消耗磁盘资源。"),

        ("九", "上传目录匿名公开访问", "上传接口/文件访问",
         "上传文件存储在 static/ 目录下，Flask 静态文件路由无任何鉴权。攻击者上传后门后可直接通过 URL 匿名访问执行。",
         "登录系统上传 shell.php 脚本，记录文件静态访问 URL。退出登录，打开无痕浏览器直接粘贴 URL 访问。",
         "WebShell 文件可被匿名访问并执行，实现远程控制，后门无鉴权可执行。",
         "将文件存储目录移出 static/ 路由范围，改为 data/uploads/ 私有目录；新增 /uploads/<filename> 路由，使用 @login_required 装饰器进行登录鉴权，仅登录用户可访问。",
         'UPLOAD_DIR = os.path.join(app.root_path, "data", "uploads")\n\n@app.route("/uploads/<filename>")\n@login_required\ndef serve_upload(filename):\n    return send_from_directory(UPLOAD_DIR, filename)',
         "文件存于 static/ 之外，无法通过 /static/ 匿名访问；访问路由要求登录，未登录用户被重定向到登录页。"),
    ]

    for num, name, loc, desc, steps, result, fix, code, principle in vulns:
        section_label = {"三": "三", "四": "四", "五": "五", "六": "六", "七": "七", "八": "八", "九": "九"}
        add_heading_styled(doc, f"{num}、漏洞 {num_to_num(num)}：{name}", level=1)
        add_separator(doc)

        add_para(doc, f"漏洞定位：{loc}", bold=True, color=(192, 57, 43))
        add_para(doc, f"【漏洞描述】{desc}", first_indent=0.74)
        add_para(doc, "【漏洞代码（修复前）】", bold=True, size=10, color=(192, 57, 43))
        add_code(doc,
            '# 修复前：原始文件名直接拼接路径\n'
            'filename = file.filename\n'
            'save_path = os.path.join("static/uploads", filename)\n'
            'file.save(save_path)'
        )
        add_para(doc, "【复现步骤】", bold=True)
        for i, s in enumerate(steps.strip().split("\n"), 1):
            add_para(doc, f"{i}. {s}", first_indent=0.74)
        add_para(doc, "【复现结果】", bold=True)
        add_para(doc, result, first_indent=0.74, color=(192, 57, 43))

        add_para(doc, "【修复方案】", bold=True)
        add_para(doc, fix, first_indent=0.74)
        add_para(doc, "【修复后代码】", bold=True, size=10, color=(39, 174, 96))
        add_code(doc, code)
        add_para(doc, f"【修复原理】{principle}", first_indent=0.74)
        doc.add_paragraph()

    doc.add_page_break()

    # ═══ 十、修复方案汇总 ═══
    add_heading_styled(doc, "十、修复方案汇总与对比", level=1); add_separator(doc)
    make_table(doc, ["安全维度", "修复前", "修复后"],
               [["文件名处理", "原始文件名直接使用", "secure_filename + UUID 重命名"],
                ["文件后缀校验", "无限制，任意后缀可上传", "白名单仅 jpg/jpeg/png/gif"],
                ["文件内容校验", "不校验，仅靠前端 Content-Type", "二进制魔数校验图片真实类型"],
                ["文件覆盖", "同名文件互相覆盖", "UUID 唯一命名，永不重复"],
                ["XSS 防护", "文件名直接渲染不转义", "Jinja2 autoescape + 后端过滤"],
                ["单文件大小", "无限制（仅全局 16MB）", "5MB 上限提前拦截"],
                ["访问鉴权", "static/ 目录匿名可访问", "data/uploads/ + @login_required 路由"],
                ["异常处理", "返回系统绝对路径", "统一错误提示，无路径泄露"]],
               col_widths=[3, 5.5, 6])

    # ═══ 十一、修复后安全验证 ═══
    add_heading_styled(doc, "十一、修复后安全验证", level=1); add_separator(doc)
    make_table(doc, ["测试用例", "预期", "实际", "结论"],
               [["正常上传 PNG 图片", "成功", "成功，UUID 文件名保存", "通过"],
                ["路径遍历 ../app.py", "拦截", "后缀白名单拒绝 .py", "通过"],
                ["上传 shell.php", "拦截", "后缀白名单拒绝 .php", "通过"],
                ["MIME 伪装 PHP→PNG", "拦截", "魔数校验失败拒绝", "通过"],
                ["XSS 文件名 <script>", "转义不执行", "显示 &lt;script&gt;", "通过"],
                ["超大文件 15MB", "拦截", "5MB 上限拒绝", "通过"],
                ["匿名访问上传文件", "重定向登录页", "302 跳转", "通过"]])

    # ═══ 十二、总结 ═══
    add_heading_styled(doc, "十二、实验总结与心得", level=1); add_separator(doc)
    add_para(doc, "通过本次文件上传漏洞手工复现与修复实验，深入理解了文件上传场景中 7 类典型安全漏洞的利用原理与修复方案。", first_indent=0.74)
    add_para(doc, "核心认知：", bold=True)
    add_para(doc, "（1）文件上传是 Web 应用最高危的功能点之一，必须对文件名、文件类型、文件大小、访问权限进行全面校验，任一环节缺失都可能导致严重安全事件。", first_indent=0.74)
    add_para(doc, "（2）永远不要信任用户提供的文件名和 Content-Type 请求头。文件名应使用 UUID 重命名，文件类型应通过二进制魔数校验。", first_indent=0.74)
    add_para(doc, "（3）上传文件的存储位置必须与静态资源目录隔离，并通过登录鉴权保护，防止后门文件被匿名访问和执行。", first_indent=0.74)

    doc.add_paragraph(); add_separator(doc)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("报告完"); set_run_font(run, size=11, color=(150, 150, 150))

    doc.save(output_path)
    print(f"报告已生成：{output_path}")


def num_to_num(s):
    m = {"三": "1", "四": "2", "五": "3", "六": "4", "七": "5", "八": "6", "九": "7"}
    return m.get(s, s)


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "文件上传漏洞复现及修复报告.docx")
    generate_report(out)
