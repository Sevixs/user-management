#!/usr/bin/env python3
"""
SQL注入漏洞手工复现及修复报告 — Word 文档生成脚本
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="微软雅黑", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return heading


def add_para(doc, text, bold=False, size=11, color=None, align=None,
             space_after=6, space_before=0, first_indent=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = Pt(20)
    if align:
        para.alignment = align
    if first_indent:
        para.paragraph_format.first_line_indent = Cm(first_indent)
    return para


def add_code_block(doc, code, indent=1.0):
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)


def add_separator(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def make_table(doc, headers, rows, col_widths=None, header_color="C0392B"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def generate_report(output_path):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(11)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ═══════════════ 封面 ═══════════════

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Flask 用户信息管理平台")
    set_run_font(run, size=26, bold=True, color=(192, 57, 43))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("SQL 注入漏洞手工复现及修复报告")
    set_run_font(run, size=20, bold=False, color=(89, 89, 89))

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━" * 40)
    set_run_font(run, size=12, color=(200, 200, 200))
    doc.add_paragraph()

    for text in [
        "实验性质：SQL 注入漏洞复现与安全加固实训",
        "目标系统：Flask 用户信息管理系统（SQLite）",
        "文档版本：V1.0 — 终审版",
        "生成日期：2026 年 7 月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=12, color=(100, 100, 100))

    doc.add_page_break()

    # ═══════════════ 目录 ═══════════════

    add_heading_styled(doc, "目  录", level=1)
    doc.add_paragraph()

    toc_items = [
        ("一、实验概述", ""),
        ("二、漏洞总览", ""),
        ("三、漏洞 1：搜索接口 UNION 联合查询注入", ""),
        ("    3.1  漏洞描述", ""),
        ("    3.2  漏洞 Payload", ""),
        ("    3.3  复现步骤", ""),
        ("    3.4  复现结果", ""),
        ("    3.5  修复方案", ""),
        ("四、漏洞 2：搜索接口 OR 万能条件拖库注入", ""),
        ("    4.1  漏洞描述", ""),
        ("    4.2  漏洞 Payload", ""),
        ("    4.3  复现步骤", ""),
        ("    4.4  复现结果", ""),
        ("    4.5  修复方案", ""),
        ("五、漏洞 3：注册接口 INSERT 注入", ""),
        ("    5.1  漏洞描述", ""),
        ("    5.2  漏洞 Payload", ""),
        ("    5.3  复现步骤", ""),
        ("    5.4  复现结果", ""),
        ("    5.5  修复方案", ""),
        ("六、修复方案汇总与对比", ""),
        ("七、修复后安全验证", ""),
        ("八、实验总结与心得", ""),
    ]
    for title_text, _ in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(f"{title_text}")
        set_run_font(run, size=11, bold=not title_text.startswith("    "))

    doc.add_page_break()

    # ═══════════════ 一、实验概述 ═══════════════

    add_heading_styled(doc, "一、实验概述", level=1)
    add_separator(doc)

    add_para(doc,
             "本次实验针对基于 Python Flask 框架开发的自定义漏洞靶场中存在的三处高危 SQL 注入漏洞，"
             "进行手工复现与安全修复。该靶场原本使用 f-string 字符串拼接方式构建 SQL 语句，"
             "用户输入完全可控，存在严重的 SQL 注入安全风险。",
             first_indent=0.74)

    add_para(doc, "本次实验涉及的三处漏洞：", bold=True, size=11)
    add_para(doc, "（1）搜索接口 UNION 联合查询注入 —— 数据读取类注入", first_indent=0.74)
    add_para(doc, "（2）搜索接口 OR 万能条件拖库注入 —— 布尔永真注入", first_indent=0.74)
    add_para(doc, "（3）注册接口 INSERT 语句注入 —— 恶意写入型注入", first_indent=0.74)

    add_para(doc,
             "实验流程包括：漏洞审计、Payload 构造、Burp Suite 手工复现、漏洞根因分析、"
             "参数化查询修复、修复后安全验证，完整覆盖 SQL 注入漏洞从发现到修复的全生命周期。",
             first_indent=0.74)

    doc.add_paragraph()

    # ═══════════════ 二、漏洞总览 ═══════════════

    add_heading_styled(doc, "二、漏洞总览", level=1)
    add_separator(doc)

    make_table(doc,
               ["编号", "漏洞名称", "漏洞类型", "危害等级", "攻击效果"],
               [
                   ["VUL-SQL-01", "搜索 UNION 联合查询注入", "数据读取", "高危",
                    "读取 users 表全部敏感数据（用户名、密码、邮箱、手机号）"],
                   ["VUL-SQL-02", "搜索 OR 万能条件拖库注入", "布尔绕过", "高危",
                    "无条件查询全表数据，一键拖库"],
                   ["VUL-SQL-03", "注册 INSERT 语句注入", "恶意写入", "高危",
                    "任意自定义新增管理员账号，后台权限接管"],
               ],
               col_widths=[2.5, 4, 2.5, 1.5, 5.5])

    add_para(doc, "漏洞根因（三处漏洞共性）：", bold=True, size=11, space_before=6)
    add_para(doc, "（1）未使用参数化预编译查询，SQL 语句结构与用户输入数据未隔离；", first_indent=0.74)
    add_para(doc, "（2）未过滤单引号、括号、SQL 注释符、逻辑运算符等危险语法字符；", first_indent=0.74)
    add_para(doc, "（3）用户可控参数直接参与 SQL 语法解析，攻击者可任意增删改查 SQL 逻辑。", first_indent=0.74)

    doc.add_page_break()

    # ═══════════════ 三、漏洞1 UNION注入 ═══════════════

    add_heading_styled(doc, "三、漏洞 1：搜索接口 UNION 联合查询注入", level=1)
    add_separator(doc)

    add_heading_styled(doc, "3.1  漏洞描述", level=2)
    add_para(doc,
             "搜索功能将前端 keyword 参数直接通过 f-string 拼接至 SELECT 查询语句，"
             "用户输入完全可控，未做任何语义隔离。攻击者可利用单引号闭合原有 LIKE 模糊查询语法，"
             "截断原有 SQL 逻辑，通过 UNION 拼接全新自定义查询语句，实现跨字段、跨数据行读取数据库任意数据。",
             first_indent=0.74)

    add_para(doc, "漏洞代码片段（修复前）：", bold=True, size=10, color=(192, 57, 43))
    add_code_block(doc,
        '# ❌ 修复前：f-string 直接拼接，存在 UNION 注入\n'
        'sql = f"SELECT id, username, email, phone FROM users "\n'
        '      f"WHERE username LIKE \'%{keyword}%\' OR email LIKE \'%{keyword}%\'"\n'
        'c.execute(sql)'
    )

    add_para(doc, "攻击特征：必须严格匹配数据表查询列数（4 列），否则 SQLite 直接报错，属于结构化严格型显错注入。", first_indent=0.74)

    add_heading_styled(doc, "3.2  漏洞 Payload", level=2)
    add_para(doc, "构造的注入 Payload 如下：", first_indent=0.74)
    add_code_block(doc, "' UNION SELECT 1,username,email,phone FROM users-- ")

    add_para(doc, "拼接后的实际 SQL 语句：", first_indent=0.74)
    add_code_block(doc,
        "SELECT id, username, email, phone FROM users\n"
        "WHERE username LIKE '%' UNION SELECT 1,username,email,phone FROM users--%'\n"
        "     OR email LIKE '%' UNION SELECT 1,username,email,phone FROM users--%'"
    )

    add_heading_styled(doc, "3.3  复现步骤（Burp Suite 手工操作）", level=2)

    steps = [
        "浏览器登录后台：使用 admin / admin123 登录，获得有效会话 Cookie。",
        "开启 Burp 代理拦截，访问搜索功能，拦截 GET /search 请求。",
        "将请求发送至 Repeater（右键 -> Send to Repeater）。",
        "修改 keyword 参数值为 UNION 注入 Payload：' UNION SELECT 1,username,email,phone FROM users--",
        "点击 Send 发送请求，观察响应结果。"
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_heading_styled(doc, "3.4  复现结果", level=2)
    add_para(doc,
             "页面成功查询并展示数据库中所有用户数据，包含 admin、alice 的用户名、邮箱、手机号等"
             "敏感信息。证明 UNION SQL 注入漏洞存在，可批量读取数据库敏感数据，属于直接数据泄露漏洞。",
             first_indent=0.74)

    add_para(doc, "危害精准描述：可查询全部用户账号、密码、邮箱、手机号等敏感数据，支持自定义数据插入回显。", bold=True, color=(192, 57, 43))

    add_heading_styled(doc, "3.5  修复方案", level=2)
    add_para(doc, "修复核心：采用参数化预编译查询，废弃 f-string 拼接。", bold=True, size=11)
    add_para(doc, "修复后代码：", bold=True, size=10, color=(39, 174, 96))
    add_code_block(doc,
        '# ✅ 修复后：使用 ? 占位符参数化查询\n'
        'like_pattern = f"%{keyword}%"\n'
        'sql = "SELECT id, username, email, phone FROM users "\n'
        '      "WHERE username LIKE ? OR email LIKE ?"\n'
        'c.execute(sql, (like_pattern, like_pattern))'
    )
    add_para(doc,
             "修复原理：通过 ? 占位符将用户输入与 SQL 语句结构隔离，SQLite 驱动自动对参数进行转义处理，"
             "用户输入中的单引号、UNION 关键字等仅作为普通字符串文本参与 LIKE 匹配，"
             "不再参与 SQL 语法解析，攻击者无法再构造联合查询语句。",
             first_indent=0.74)

    doc.add_page_break()

    # ═══════════════ 四、漏洞2 OR注入 ═══════════════

    add_heading_styled(doc, "四、漏洞 2：搜索接口 OR 万能条件拖库注入", level=1)
    add_separator(doc)

    add_heading_styled(doc, "4.1  漏洞描述", level=2)
    add_para(doc,
             "搜索条件 WHERE 子句完全由用户输入拼接，未做逻辑校验。攻击者通过构造永真逻辑条件，"
             "使整条 WHERE 查询条件恒成立，完全绕过业务搜索过滤逻辑，无条件查询全表数据。"
             "无需匹配列数、无需联合查询，仅篡改查询逻辑，属于最简布尔条件绕过注入。",
             first_indent=0.74)

    add_heading_styled(doc, "4.2  漏洞 Payload", level=2)
    add_code_block(doc, "' OR '1'='1")

    add_para(doc, "拼接后的实际 SQL 语句：", first_indent=0.74)
    add_code_block(doc,
        "SELECT id, username, email, phone FROM users\n"
        "WHERE username LIKE '%' OR '1'='1%'\n"
        "     OR email LIKE '%' OR '1'='1%'"
    )

    add_heading_styled(doc, "4.3  复现步骤（Burp Suite 手工操作）", level=2)
    steps = [
        "浏览器登录后台：admin / admin123，获得有效会话。",
        "Burp 拦截搜索请求，送入 Repeater。",
        "修改 keyword 参数值为 OR 永真 Payload：' OR '1'='1",
        "发送请求查看响应。"
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_heading_styled(doc, "4.4  复现结果", level=2)
    add_para(doc,
             "WHERE 条件永久成立（'1'='1 恒为真），忽略所有搜索过滤规则，"
             "页面返回数据库全部用户数据，实现完整拖库效果。",
             first_indent=0.74)

    add_para(doc, "危害精准描述：一键拖库，返回 users 表所有用户数据，是危害最高、利用门槛最低的信息泄露漏洞。", bold=True, color=(192, 57, 43))

    add_heading_styled(doc, "4.5  修复方案", level=2)
    add_para(doc, "修复核心：参数化查询隔离用户输入 + SQL 危险字符拦截。", bold=True, size=11)
    add_para(doc,
             "与 UNION 注入采用相同的参数化查询修复方案。由于用户输入通过 ? 占位符绑定，输入内容不再参与 SQL 语法结构，"
             "OR '1'='1 仅作为普通字符串进行 LIKE 模糊匹配，无法篡改 WHERE 子句逻辑。",
             first_indent=0.74)
    add_para(doc,
             "此外，新增 validate_sql_input() 函数对输入进行 SQL 危险字符黑名单拦截，"
             "单引号、分号、注释符等字符被直接拒绝，从输入层增加一道防线。",
             first_indent=0.74)

    add_para(doc, "新增校验函数：", bold=True, size=10, color=(39, 174, 96))
    add_code_block(doc,
        'def validate_sql_input(text, field_type="text"):\n'
        '    """校验用户输入，拦截 SQL 注入危险字符，并做格式白名单校验。"""\n'
        '    # SQL 危险字符黑名单\n'
        '    dangerous = re.search(r"[\\\'\\"\\;\\-\\-/*\\\\\\`\\(\\)]", text)\n'
        '    if dangerous:\n'
        '        return False, "输入包含非法 SQL 字符"\n\n'
        '    # 格式白名单校验\n'
        '    if field_type == "username":\n'
        '        if not re.match(r"^[a-zA-Z0-9_@.\\-]+$", text):\n'
        '            return False, "用户名只能包含字母、数字、下划线、@、点、横线"\n'
        '    ...\n'
        '    return True, ""'
    )

    doc.add_page_break()

    # ═══════════════ 五、漏洞3 INSERT注入 ═══════════════

    add_heading_styled(doc, "五、漏洞 3：注册接口 INSERT 语句注入", level=1)
    add_separator(doc)

    add_heading_styled(doc, "5.1  漏洞描述", level=2)
    add_para(doc,
             "注册页面将 username 参数直接通过 f-string 拼接至 INSERT 语句，用户输入可篡改 VALUES 字段结构。"
             "攻击者利用单引号闭合当前字段，手动补全后续所有数据库字段、闭合括号，"
             "利用 SQL 行注释 -- 废弃后端原本的表单参数（password/email/phone），"
             "单字段控制整行数据库写入内容。",
             first_indent=0.74)

    add_para(doc, "漏洞代码片段（修复前）：", bold=True, size=10, color=(192, 57, 43))
    add_code_block(doc,
        "# ❌ 修复前：f-string 拼接 INSERT 语句\n"
        'sql = f"INSERT INTO users (username, password, email, phone) "\n'
        '      f"VALUES (\'{username}\', \'{password}\', \'{email}\', \'{phone}\')"\n'
        "c.execute(sql)"
    )

    add_heading_styled(doc, "5.2  漏洞 Payload", level=2)
    add_para(doc, "在 username 字段中注入：", first_indent=0.74)
    add_code_block(doc, "hacker6', 'pass123', 'hack@x.com', '13888888888') -- ")

    add_para(doc, "拼接后的实际 SQL 语句：", first_indent=0.74)
    add_code_block(doc,
        "INSERT INTO users (username, password, email, phone)\n"
        "VALUES ('hacker6', 'pass123', 'hack@x.com', '13888888888') -- ',\n"
        "       'xxx', 'yyy', 'zzz')"
    )

    add_heading_styled(doc, "5.3  复现步骤（Burp Suite 手工操作）", level=2)
    steps = [
        "浏览器进入注册页面，随意输入内容，点击注册提交。",
        "Burp 拦截 POST /register 请求。",
        "发送到 Repeater。",
        "将 username 值替换为 Payload：hacker6', 'pass123', 'hack@x.com', '13888888888') --",
        "保持 password、email、phone 为任意内容，点击 Send。",
        "访问登录页，使用 用户名：hacker6、密码：pass123 登录验证。"
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_heading_styled(doc, "5.4  复现结果", level=2)
    add_para(doc,
             "SQL 语句被成功篡改，后端自动插入攻击者自定义的账号 hacker6，"
             "密码为 pass123。使用该恶意账号可正常登录系统，实现后台权限接管。",
             first_indent=0.74)

    add_para(doc, "危害精准描述：可任意自定义用户名、密码、邮箱、手机号，强行写入恶意管理员账号，实现后台权限接管、恶意用户批量植入。", bold=True, color=(192, 57, 43))

    add_heading_styled(doc, "5.5  修复方案", level=2)
    add_para(doc, "修复核心：表单字段严格参数绑定 + 输入格式白名单校验。", bold=True, size=11)
    add_para(doc, "修复后代码：", bold=True, size=10, color=(39, 174, 96))
    add_code_block(doc,
        '# ✅ 修复后：VALUES 使用 ? 占位符参数化\n'
        'sql = "INSERT INTO users (username, password, email, phone) "\n'
        '      "VALUES (?, ?, ?, ?)"\n'
        'c.execute(sql, (username, password, email, phone))'
    )
    add_para(doc,
             "修复原理：通过 ? 占位符绑定所有表单字段，用户输入不再参与 INSERT 语句的结构解析。"
             "username 字段中的单引号、括号、注释符仅作为字符串值的一部分被插入数据库，"
             "无法闭合 VALUES 语法结构或废弃其他字段参数。同时用户名格式白名单限制仅允许"
             "字母、数字、下划线等合法字符，从输入层面进一步压缩攻击面。",
             first_indent=0.74)

    doc.add_page_break()

    # ═══════════════ 六、修复方案汇总 ═══════════════

    add_heading_styled(doc, "六、修复方案汇总与对比", level=1)
    add_separator(doc)

    make_table(doc,
               ["漏洞", "原代码（有漏洞）", "修复后代码", "修复技术"],
               [
                   ["搜索 UNION 注入",
                    "f\"...LIKE '%{keyword}%'...\"",
                    "...LIKE ? ... \nc.execute(sql, (like_pattern, like_pattern))",
                    "参数化预编译查询"],
                   ["搜索 OR 布尔注入",
                    "WHERE username LIKE '%{keyword}%'",
                    "WHERE username LIKE ? \n输入不参与语法结构",
                    "参数化 + 危险字符过滤"],
                   ["注册 INSERT 注入",
                    "f\"...VALUES ('{username}', ...)\"",
                    "VALUES (?, ?, ?, ?) \nc.execute(sql, (u, p, e, ph))",
                    "参数化预编译 + 白名单校验"],
               ],
               col_widths=[2.5, 4, 5, 3.5])

    make_table(doc,
               ["安全维度", "修复前", "修复后"],
               [
                   ["SQL 构建方式", "f-string 字符串拼接", "? 占位符参数化预编译"],
                   ["危险字符过滤", "无", "黑名单拦截 8 类 SQL 特殊字符"],
                   ["输入格式校验", "无", "白名单正则 + 长度限制"],
                   ["单引号处理", "直接拼接，可闭合语句", "占位符绑定，自动转义"],
                   ["UNION 注入防御", "可执行 UNION 查询", "输入不参与语法解析，无法构造"],
                   ["OR 永真防御", "可篡改 WHERE 逻辑", "输入仅作 LIKE 文本匹配"],
                   ["INSERT 篡改防御", "可闭合 VALUES 注入", "字段参数绑定，无法篡改结构"],
                   ["错误信息处理", "直接回显数据库异常", "统一错误提示，不暴露敏感信息"],
               ],
               col_widths=[3.5, 5.5, 6])

    doc.add_page_break()

    # ═══════════════ 七、修复后验证 ═══════════════

    add_heading_styled(doc, "七、修复后安全验证", level=1)
    add_separator(doc)

    add_para(doc, "修复完成后，对全部三处漏洞进行回归测试，验证结果如下：", first_indent=0.74)

    make_table(doc,
               ["测试用例", "预期结果", "实际结果", "结论"],
               [
                   ["正常搜索 'admin'", "返回 admin 用户信息", "返回 1 条结果", "通过"],
                   ["搜索 UNION 注入\n' UNION SELECT ...", "拦截或拒绝执行", '显示「输入包含非法 SQL 字符」', "通过"],
                   ["搜索 OR 永真\n' OR '1'='1", "拦截或拒绝执行", '显示「输入包含非法 SQL 字符」', "通过"],
                   ["正常注册合法用户", "注册成功，可登录", "注册成功，正常登录", "通过"],
                   ["注册 INSERT 注入\nusername 含闭合语法", "拦截，不写入数据", '显示「输入包含非法 SQL 字符」', "通过"],
                   ["注册重复用户名", "提示已存在", '显示「用户名已存在」', "通过"],
                   ["搜索无结果关键词", '显示「无搜索结果」', '显示「无搜索结果」', "通过"],
               ],
               col_widths=[3.5, 3.5, 4, 1.5])

    add_para(doc, "验证结论：", bold=True, size=11, space_before=8)
    add_para(doc,
             "全部三处 SQL 注入漏洞已成功修复。正常业务功能完全保留，注入 Payload 全部被拦截。"
             "修复后系统满足以下安全要求：\n"
             "（1）无字符串拼接 SQL，全部使用参数化预编译查询；\n"
             "（2）SQL 危险字符被黑名单拦截，无法传入 SQL 解析层；\n"
             "（3）输入格式白名单校验限制攻击者自由度；\n"
             "（4）错误信息统一化，不暴露数据库结构信息。",
             first_indent=0.74)

    doc.add_page_break()

    # ═══════════════ 八、总结 ═══════════════

    add_heading_styled(doc, "八、实验总结与心得", level=1)
    add_separator(doc)

    add_para(doc,
             "通过本次 SQL 注入漏洞手工复现与修复实验，深入理解了三类典型 SQL 注入攻击的利用原理、"
             "攻击链路与修复方案。",
             first_indent=0.74)

    add_para(doc, "核心认知：", bold=True, size=11)
    add_para(doc,
             "（1）SQL 注入的本质是「数据与代码未分离」。用户输入被当作 SQL 代码解析执行，"
             "而非仅作为数据使用。参数化预编译查询通过占位符机制将数据与语法结构严格隔离，"
             "是防御 SQL 注入的根本手段。",
             first_indent=0.74)
    add_para(doc,
             "（2）f-string 拼接 SQL 是最高危的编码习惯之一。即使开发者自认为做了过滤，"
             "只要存在字符串拼接且未使用占位符，就存在被绕过的风险。",
             first_indent=0.74)
    add_para(doc,
             "（3）纵深防御优于单点防御。本次修复结合了参数化查询（根本修复）+ 危险字符黑名单"
             "（输入层过滤）+ 格式白名单（语义层限制），构建了多层防护体系。",
             first_indent=0.74)
    add_para(doc,
             "（4）SQL 注入漏洞的修复必须在保证业务功能的前提下进行。修复完成后进行完整的"
             "功能回归测试与安全验证测试，确保正常功能可用、攻击向量被拦截。",
             first_indent=0.74)

    add_para(doc,
             "通过本次实验，深刻认识到任何将用户输入直接拼接入 SQL 语句的行为都是不可接受的"
             "安全红线。在今后的开发中，所有数据库操作必须使用参数化预编译查询，"
             "同时配合输入校验、最小权限原则等安全措施，构建安全的 Web 应用。",
             first_indent=0.74)

    # 结尾
    doc.add_paragraph()
    add_separator(doc)

    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run("— 报告完 —")
    set_run_font(run, size=11, color=(150, 150, 150))

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "本报告仅供安全教学与技术交流使用。报告中涉及的漏洞代码已全部修复，\n"
        "请勿将存在漏洞的代码版本部署到任何生产环境。"
    )
    set_run_font(run, size=9, color=(180, 180, 180))

    doc.save(output_path)
    print(f"报告已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "SQL注入漏洞手工复现及修复报告.docx")
    generate_report(out)
