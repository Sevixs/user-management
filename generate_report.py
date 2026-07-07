#!/usr/bin/env python3
"""
安全漏洞检测与修复报告 — Word 文档生成脚本
生成面向开发团队 / 安全审计的专业 .docx 报告。
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ──────────────────────────────────────────────────────
# 工具函数
# ──────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """设置表格单元格底色。"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="微软雅黑", size=None, bold=False, color=None):
    """统一设置 run 字体属性。"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题。"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12,
                     bold=True)
    return heading


def add_para(doc, text, bold=False, size=11, color=None, align=None,
             space_after=6, space_before=0, first_indent=None):
    """添加格式化段落。"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = Pt(20)
    if align:
        para.alignment = align
    if first_indent:
        para.paragraph_format.first_line_indent = Cm(first_indent)
    return para


def add_bullet(doc, text, level=0, size=11):
    """添加无序列表项。"""
    para = doc.add_paragraph(text, style="List Bullet")
    para.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = Pt(19)
    for run in para.runs:
        set_run_font(run, size=size)
    return para


def add_rich_para(doc, parts, space_after=6, first_indent=None):
    """添加富文本段落。parts = [(text, bold, color), ...]"""
    para = doc.add_paragraph()
    for text, bold, color in parts:
        run = para.add_run(text)
        set_run_font(run, size=11, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = Pt(20)
    if first_indent:
        para.paragraph_format.first_line_indent = Cm(first_indent)
    return para


def make_table(doc, headers, rows, col_widths=None, header_color="2F5496"):
    """创建格式化表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后间距
    return table


def add_separator(doc):
    """添加水平分隔线。"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


# ──────────────────────────────────────────────────────
# 主文档生成
# ──────────────────────────────────────────────────────

def generate_report(output_path):
    doc = Document()

    # ── 全局样式 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(11)

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ══════════════════════════════════════════════════
    # 封面
    # ══════════════════════════════════════════════════

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Flask 用户信息管理平台")
    set_run_font(run, size=26, bold=True, color=(47, 84, 150))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("漏洞检测与安全修复报告")
    set_run_font(run, size=22, bold=False, color=(89, 89, 89))

    doc.add_paragraph()

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━" * 40)
    set_run_font(run, size=12, color=(200, 200, 200))

    doc.add_paragraph()

    for text in [
        "实验性质：Web 安全漏洞审计与加固实训",
        "目标系统：Flask 用户信息管理系统",
        "文档版本：V1.0 — 终审版",
        "生成日期：2026 年 7 月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=12, color=(100, 100, 100))

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 目录页
    # ══════════════════════════════════════════════════

    add_heading_styled(doc, "目  录", level=1)
    doc.add_paragraph()

    toc_items = [
        ("一、实验概述", "3"),
        ("二、实验环境", "3"),
        ("三、原始系统安全漏洞审计结果", "4"),
        ("    3.1  高危漏洞 1：用户密码明文存储与明文校验", "4"),
        ("    3.2  高危漏洞 2：前端源代码注释泄露管理员账号", "5"),
        ("    3.3  高危漏洞 3：页面回显完整敏感用户信息", "5"),
        ("    3.4  高危漏洞 4：无 CSRF 跨站请求伪造防护", "6"),
        ("    3.5  中危漏洞 5：用户输入未过滤 — XSS 注入风险", "6"),
        ("    3.6  中危漏洞 6：服务端调试模式常开", "7"),
        ("    3.7  低危漏洞 7：固定弱密钥导致会话可伪造", "7"),
        ("四、漏洞修复原理与具体实施步骤", "8"),
        ("    4.1  密码安全重构（漏洞 1）", "8"),
        ("    4.2  信息泄露修复（漏洞 2、3）", "9"),
        ("    4.3  CSRF 防护加固（漏洞 4）", "9"),
        ("    4.4  XSS 跨站脚本修复（漏洞 5）", "10"),
        ("    4.5  服务端安全配置整改（漏洞 6、7）", "10"),
        ("    4.6  路由权限鉴权补全", "11"),
        ("五、修复后功能与安全性测试", "11"),
        ("六、修复前后安全对比分析", "12"),
        ("七、安全优化总结与后续改进方案", "13"),
        ("八、实验心得", "14"),
    ]
    for title_text, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(f"{title_text}")
        set_run_font(run, size=11, bold=not title_text.startswith("    "))
        run2 = p.add_run(f"  ·  {page}")
        set_run_font(run2, size=10, color=(150, 150, 150))

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 一、实验概述
    # ══════════════════════════════════════════════════

    add_heading_styled(doc, "一、实验概述", level=1)
    add_separator(doc)

    add_para(doc,
             "本次实验对象为基于 Python Flask 框架开发的简易用户信息管理系统，系统具备账号登录、"
             "会话保持、个人信息展示、用户登出等基础 Web 业务功能。初始版本代码为漏洞演示版本，"
             "存在多处典型 Web 安全缺陷，包含明文密码存储、敏感信息泄露、缺乏请求防护、输入未过滤、"
             "配置不安全等问题。",
             first_indent=0.74)

    add_para(doc,
             "本次实验通过代码审计、漏洞复现、安全加固、功能回归测试完整流程，完成系统全部高危、"
             "中危漏洞的修复，在不破坏原有业务功能的前提下提升系统整体安全等级，完成本次安全整改实训。",
             first_indent=0.74)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 二、实验环境
    # ══════════════════════════════════════════════════

    add_heading_styled(doc, "二、实验环境", level=1)
    add_separator(doc)

    make_table(doc,
               ["项目", "规格 / 版本"],
               [
                   ["开发语言", "Python 3.10+"],
                   ["开发框架", "Flask 3.1 + Werkzeug 3.1"],
                   ["模板引擎", "Jinja2（Flask 内置）"],
                   ["前端技术", "HTML5 + CSS3（Flexbox 布局）"],
                   ["密码加密", "Werkzeug Security（scrypt / pbkdf2:sha256）"],
                   ["CSRF 实现", "secrets 模块 + before_request 全局校验"],
                   ["运行地址", "http://0.0.0.0:5000"],
                   ["测试方式", "源码审计 / 人工渗透测试 / 功能回归测试"],
               ],
               col_widths=[4, 12])

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 三、原始系统安全漏洞审计结果
    # ══════════════════════════════════════════════════

    add_heading_styled(doc, "三、原始系统安全漏洞审计结果", level=1)
    add_separator(doc)

    add_para(doc,
             "通过逐行代码审计与渗透测试，本系统初始版本共发现 7 处安全漏洞，包含高危 4 项、"
             "中危 2 项、低危 1 项。漏洞总览如下表所示：",
             first_indent=0.74)

    # 漏洞总览表
    make_table(doc,
               ["编号", "漏洞名称", "严重程度", "影响范围"],
               [
                   ["VUL-01", "用户密码明文存储与明文校验", "高危", "全部用户账号"],
                   ["VUL-02", "前端源代码注释泄露管理员账号", "高危", "管理员账号"],
                   ["VUL-03", "页面回显完整敏感用户信息", "高危", "全部用户隐私数据"],
                   ["VUL-04", "无 CSRF 跨站请求伪造防护", "高危", "全部表单提交"],
                   ["VUL-05", "用户输入未过滤 — XSS 注入风险", "中危", "全部页面"],
                   ["VUL-06", "服务端调试模式常开，报错泄露源码", "中危", "系统架构信息"],
                   ["VUL-07", "固定弱密钥导致会话可伪造", "低危", "全部会话"],
               ],
               col_widths=[2, 7, 2, 4])

    # ── 3.1 高危漏洞1 ──
    add_heading_styled(doc, "3.1  高危漏洞 1：用户密码明文存储与明文校验", level=2)

    add_para(doc, "【漏洞现象】", bold=True, size=11)
    add_para(doc,
             "系统使用字典变量 USERS 存储用户数据，所有用户密码以明文硬编码保存在 Python 源码中。"
             "登录逻辑采用==字符串直接比对方式验证身份，未做任何加密处理。",
             first_indent=0.74)
    add_para(doc, "漏洞代码片段：", bold=True, size=10, color=(180, 60, 60))
    code_text = (
        'USERS = {\n'
        '    "admin": {\n'
        '        "username": "admin",\n'
        '        "password": "admin123",   # 明文存储\n'
        '        ...\n'
        '    },\n'
        '    ...\n'
        '}\n'
        '\n'
        '# 登录时直接明文比对\n'
        'if username in USERS and USERS[username]["password"] == password:'
    )
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    add_para(doc, "【安全风险】", bold=True, size=11)
    add_para(doc,
             "一旦源码泄露、页面数据被抓取或开发人员终端被入侵，全部用户账号密码直接暴露。"
             "攻击者可利用泄露的账号密码直接登录系统，造成权限失控、数据泄露等严重后果。",
             first_indent=0.74)
    add_para(doc, "【风险等级】高危  CVSS 3.1 Score: 9.8（Critical）", bold=True,
             color=(200, 50, 50))

    # ── 3.2 高危漏洞2 ──
    add_heading_styled(doc, "3.2  高危漏洞 2：前端源代码注释泄露管理员账号", level=2)

    add_para(doc, "【漏洞现象】", bold=True, size=11)
    add_para(doc,
             "登录页面 login.html 源码中存在明文 HTML 注释，直接标注默认管理员账号与密码信息。"
             "任意用户通过浏览器查看网页源代码即可获取最高权限账号。",
             first_indent=0.74)
    add_para(doc, "漏洞代码片段：", bold=True, size=10, color=(180, 60, 60))
    p = doc.add_paragraph()
    run = p.add_run(
        '<!-- 调试信息 - 默认管理员账号 用户名: admin 密码: admin123 -->'
    )
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)

    add_para(doc, "【安全风险】", bold=True, size=11)
    add_para(doc,
             "属于直接敏感信息泄露，无任何防护措施即可被未授权人员获取系统核心账号，"
             "极大提升被入侵风险。攻击者可利用泄露的管理员账号直接接管系统。",
             first_indent=0.74)
    add_para(doc, "【风险等级】高危  CVSS 3.1 Score: 8.6（High）", bold=True,
             color=(200, 50, 50))

    # ── 3.3 高危漏洞3 ──
    add_heading_styled(doc, "3.3  高危漏洞 3：页面回显完整敏感用户信息", level=2)

    add_para(doc, "【漏洞现象】", bold=True, size=11)
    add_para(doc,
             "登录成功后首页 index.html 直接渲染并展示用户密码明文、手机号、余额、邮箱等"
             "全部隐私字段，无任何脱敏处理，前端浏览器缓存 / 历史记录均可被窃取。",
             first_indent=0.74)
    add_para(doc, "漏洞代码片段：", bold=True, size=10, color=(180, 60, 60))
    p = doc.add_paragraph()
    run = p.add_run(
        '<li><span>密码：</span>{{ user.password }}</li>\n'
        '<li><span>手机：</span>{{ user.phone }}</li>\n'
        '<li><span>余额：</span>{{ user.balance }}</li>'
    )
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)

    add_para(doc, "【安全风险】", bold=True, size=11)
    add_para(doc,
             "违反数据最小化展示原则，造成用户隐私数据大面积泄露。"
             "密码明文回显意味着只要用户登录过一次，其密码就会留存于浏览器历史、"
             "代理缓存、开发者工具日志等多个可被利用的位置。",
             first_indent=0.74)
    add_para(doc, "【风险等级】高危  CVSS 3.1 Score: 7.5（High）", bold=True,
             color=(200, 50, 50))

    # ── 3.4 高危漏洞4 ──
    add_heading_styled(doc, "3.4  高危漏洞 4：无 CSRF 跨站请求伪造防护", level=2)

    add_para(doc, "【漏洞现象】", bold=True, size=11)
    add_para(doc,
             "系统登录 POST 表单未配置 CSRF 令牌校验机制，所有跨站请求均可直接提交至服务器。"
             "攻击者可构造恶意页面，诱导已登录用户执行非预期的操作。",
             first_indent=0.74)

    add_para(doc, "【安全风险】", bold=True, size=11)
    add_para(doc,
             "攻击者可构造恶意站点嵌入隐藏表单或自动提交脚本，当已登录用户访问该恶意站点时，"
             "浏览器会自动携带目标站点的 Cookie 发起 POST 请求，实现 CSRF 劫持攻击。",
             first_indent=0.74)
    add_para(doc, "【风险等级】高危  CVSS 3.1 Score: 8.8（High）", bold=True,
             color=(200, 50, 50))

    # ── 3.5 中危漏洞5 ──
    add_heading_styled(doc, "3.5  中危漏洞 5：用户输入未过滤 — XSS 注入风险", level=2)

    add_para(doc, "【漏洞现象】", bold=True, size=11)
    add_para(doc,
             "登录接口直接接收前端表单参数 username / password，未对特殊字符、HTML 标签、"
             "JavaScript 脚本关键字进行过滤与转义，直接传递至模板渲染。",
             first_indent=0.74)

    add_para(doc, "【安全风险】", bold=True, size=11)
    add_para(doc,
             "攻击者可植入 JavaScript 恶意代码，实现弹窗劫持、Cookie 窃取、"
             "页面篡改、跳转钓鱼页面等 XSS 攻击。虽然 Jinja2 模板引擎默认转义，"
             "但在复杂场景下或使用 |safe 过滤器时存在被绕过风险。",
             first_indent=0.74)
    add_para(doc, "【风险等级】中危  CVSS 3.1 Score: 6.1（Medium）", bold=True,
             color=(200, 130, 0))

    # ── 3.6 中危漏洞6 ──
    add_heading_styled(doc, "3.6  中危漏洞 6：服务端调试模式常开", level=2)

    add_para(doc, "【漏洞现象】", bold=True, size=11)
    add_para(doc,
             "程序启动参数固定设置 debug=True，生产模式开启 Flask 调试模式。"
             "应用发生异常时会展示 Werkzeug 调试器页面，包含完整调用栈与交互式 Shell。",
             first_indent=0.74)
    add_para(doc, "漏洞代码片段：", bold=True, size=10, color=(180, 60, 60))
    p = doc.add_paragraph()
    run = p.add_run('app.run(debug=True, host="0.0.0.0", port=5000)')
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)

    add_para(doc, "【安全风险】", bold=True, size=11)
    add_para(doc,
             "系统报错时泄露源码路径、项目结构、环境变量、服务器操作系统等信息，"
             "Werkzeug 调试器甚至允许攻击者在服务器上执行任意 Python 代码（RCE）。",
             first_indent=0.74)
    add_para(doc, "【风险等级】中危  CVSS 3.1 Score: 5.3（Medium）", bold=True,
             color=(200, 130, 0))

    # ── 3.7 低危漏洞7 ──
    add_heading_styled(doc, "3.7  低危漏洞 7：固定弱密钥导致会话可伪造", level=2)

    add_para(doc, "【漏洞现象】", bold=True, size=11)
    add_para(doc,
             "系统密钥硬编码为 dev-key-2025，密钥强度极低、可被猜测。"
             "所有会话的签名均使用该固定密钥，且无会话超时机制。",
             first_indent=0.74)
    add_para(doc, "漏洞代码片段：", bold=True, size=10, color=(180, 60, 60))
    p = doc.add_paragraph()
    run = p.add_run('app.secret_key = "dev-key-2025"')
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)

    add_para(doc, "【安全风险】", bold=True, size=11)
    add_para(doc,
             "已知密钥可被攻击者用于伪造任意用户的 Session Cookie，实现身份冒充。"
             "同时无会话超时机制意味着一次登录永久有效，增大了会话劫持的攻击窗口。",
             first_indent=0.74)
    add_para(doc, "【风险等级】低危  CVSS 3.1 Score: 3.7（Low）", bold=True,
             color=(0, 120, 200))

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 四、漏洞修复原理与具体实施步骤
    # ══════════════════════════════════════════════════

    add_heading_styled(doc, "四、漏洞修复原理与具体实施步骤", level=1)
    add_separator(doc)

    add_para(doc,
             "本次修复遵循「保留全部业务功能、逐项定点整改、由高危到低危」的修复原则。"
             "以下逐条说明每项漏洞的修复方案、原理与代码位置。",
             first_indent=0.74)

    # ── 4.1 ──
    add_heading_styled(doc, "4.1  密码安全重构（漏洞 1）", level=2)

    add_para(doc, "修复方案：", bold=True, size=11)
    add_para(doc,
             "采用 Werkzeug.security 模块的 generate_password_hash() 函数，使用 scrypt "
             "迭代算法（默认配置 32768 轮、8 并行、1 线程）对密码进行不可逆哈希加密。"
             "USERS 字典中仅存储哈希值。登录验证时使用 check_password_hash() 进行安全比对，"
             "原始密码仅在 _build_user_db() 函数中出现，运行后内存中仅保留哈希密文。",
             first_indent=0.74)

    add_para(doc, "修复原理：", bold=True, size=11)
    add_para(doc,
             "PBKDF2 / scrypt 算法为密码加盐后迭代哈希，即使哈希值泄露，攻击者也无法逆推出"
             "原始密码，必须通过暴力枚举或彩虹表方式进行破解。选择 scrypt 算法因其同时消耗"
             "CPU 与内存资源，大幅提升 GPU 并行破解的硬件成本。",
             first_indent=0.74)

    add_para(doc, "核心代码（app.py）：", bold=True, size=10, color=(47, 84, 150))
    code = (
        'from werkzeug.security import generate_password_hash, check_password_hash\n\n'
        'def _build_user_db():\n'
        '    raw = {"admin": {"password": "admin123", ...}}\n'
        '    db = {}\n'
        '    for uid, info in raw.items():\n'
        '        record = info.copy()\n'
        '        record["password"] = generate_password_hash(record["password"])\n'
        '        db[uid] = record\n'
        '    return db\n\n'
        '# 登录验证\n'
        'if user_record and check_password_hash(user_record["password"], password):'
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)

    # ── 4.2 ──
    add_heading_styled(doc, "4.2  信息泄露修复（漏洞 2、3）", level=2)

    add_para(doc, "修复方案：", bold=True, size=11)

    add_para(doc, "（1）删除前端注释泄露", bold=True, size=10)
    add_para(doc,
             "清理 login.html 页面全部包含账号密码的 HTML 注释，删除所有开发阶段遗留的"
             "敏感调试信息，保证前端源码无任何可利用的泄露点。",
             first_indent=0.74)

    add_para(doc, "（2）敏感数据脱敏展示", bold=True, size=10)
    add_para(doc,
             "新增 sanitize_user_info() 函数，在将用户信息传入模板前进行数据脱敏："
             "彻底移除 password 字段（永不传递到前端）；手机号中间四位替换为 ****，"
             "如 138****8000；余额格式化为 ¥99,999.00 货币显示。",
             first_indent=0.74)

    add_para(doc, "核心代码（app.py）：", bold=True, size=10, color=(47, 84, 150))
    code = (
        'def sanitize_user_info(user_info):\n'
        '    """构造可供模板安全使用的用户信息字典。"""\n'
        '    if not user_info:\n'
        '        return None\n'
        '    return {\n'
        '        "username": user_info.get("username", ""),\n'
        '        "role": user_info.get("role", ""),\n'
        '        "email": user_info.get("email", ""),\n'
        '        "phone": user_info["phone"][:3] + "****" + user_info["phone"][-4:],\n'
        '        "balance": "¥{:,.2f}".format(float(user_info["balance"])),\n'
        '        # password 字段被排除\n'
        '    }'
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)

    # ── 4.3 ──
    add_heading_styled(doc, "4.3  CSRF 防护加固（漏洞 4）", level=2)

    add_para(doc, "修复方案：", bold=True, size=11)
    add_para(doc,
             "引入 Python 标准库 secrets 模块实现的 CSRF 令牌机制，"
             "不依赖第三方扩展。通过 @app.before_request 全局拦截所有 POST 请求，"
             "校验表单中的 csrf_token 字段是否与会话中存储的令牌一致。"
             "令牌通过 @app.context_processor 注入所有模板，模板中渲染为隐藏字段。",
             first_indent=0.74)

    add_para(doc, "修复原理：", bold=True, size=11)
    add_para(doc,
             "CSRF Token 是一种同步令牌模式（Synchronizer Token Pattern）：服务器为每个会话"
             "生成唯一的随机令牌，嵌入表单中提交。由于攻击者无法获取目标用户的会话令牌，"
             "因此跨站构造的请求无法通过校验。本实现使用 secrets.token_hex(32) 生成 256 位"
             "密码学安全随机数，保证令牌不可预测。",
             first_indent=0.74)

    add_para(doc, "核心代码（app.py & login.html）：", bold=True, size=10,
             color=(47, 84, 150))
    code = (
        '# app.py — CSRF 校验\n'
        '@app.before_request\n'
        'def _csrf_protect():\n'
        '    if request.method == "POST" and request.endpoint != "static":\n'
        '        token = request.form.get("csrf_token")\n'
        '        stored = session.get("csrf_token")\n'
        '        if not token or not stored or token != stored:\n'
        '            abort(400, "CSRF token 缺失或无效")\n\n'
        '@app.context_processor\n'
        'def _inject_csrf_token():\n'
        '    if "csrf_token" not in session:\n'
        '        session["csrf_token"] = secrets.token_hex(32)\n'
        '    return {"csrf_token": session["csrf_token"]}\n\n'
        '<!-- login.html — 表单内添加 -->\n'
        '<input type="hidden" name="csrf_token" value="{{ csrf_token }}">'
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)

    # ── 4.4 ──
    add_heading_styled(doc, "4.4  XSS 跨站脚本修复（漏洞 5）", level=2)

    add_para(doc, "修复方案：", bold=True, size=11)
    add_para(doc,
             "（1）编写 sanitize_input() 函数，对用户输入的用户名字段进行严格过滤："
             "移除所有 HTML 标签（<...>）；移除 javascript: 伪协议关键字；"
             "去除首尾空白字符。\n"
             "（2）在登录表单 input 标签中增加 maxlength 属性限制输入长度。\n"
             "（3）利用 Jinja2 模板引擎默认启用的 autoescape 机制，{{ ... }} 输出时"
             "自动转义 HTML 特殊字符为实体编码。",
             first_indent=0.74)

    add_para(doc, "核心代码（app.py）：", bold=True, size=10, color=(47, 84, 150))
    code = (
        'def sanitize_input(text):\n'
        '    """过滤用户输入，移除 HTML 标签和危险关键字。"""\n'
        '    if not text:\n'
        '        return ""\n'
        '    text = str(text).strip()\n'
        '    text = re.sub(r"<[^>]*>", "", text)\n'
        '    text = re.sub(r"javascript\\s*:", "", text, flags=re.IGNORECASE)\n'
        '    return text\n\n'
        '# 使用\n'
        'username = sanitize_input(request.form.get("username", ""))'
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)

    # ── 4.5 ──
    add_heading_styled(doc, "4.5  服务端安全配置整改（漏洞 6、7）", level=2)

    add_para(doc, "修复方案：", bold=True, size=11)
    add_para(doc,
             "（1）Debug 模式安全化：将 app.run(debug=True) 修改为由环境变量 "
             "FLASK_DEBUG 控制，默认关闭调试模式，生产环境通过 export FLASK_DEBUG=true "
             "按需开启。\n"
             "（2）密钥安全化：废弃硬编码 dev-key-2025，优先读取环境变量 SECRET_KEY，"
             "未设置时自动使用 secrets.token_hex(32) 生成 256 位随机密钥。\n"
             "（3）会话安全增强：设置 PERMANENT_SESSION_LIFETIME = 2 小时超时；"
             "SESSION_COOKIE_HTTPONLY = True 禁止 JavaScript 读取 Cookie；"
             "SESSION_COOKIE_SAMESITE = Lax 禁止跨站携带 Cookie。",
             first_indent=0.74)

    add_para(doc, "核心代码（app.py）：", bold=True, size=10, color=(47, 84, 150))
    code = (
        'import os\n'
        'import secrets\n'
        'from datetime import timedelta\n\n'
        'app.config.update(\n'
        '    SECRET_KEY=os.environ.get("SECRET_KEY", secrets.token_hex(32)),\n'
        '    PERMANENT_SESSION_LIFETIME=timedelta(hours=2),\n'
        '    SESSION_COOKIE_HTTPONLY=True,\n'
        '    SESSION_COOKIE_SAMESITE="Lax",\n'
        '    SESSION_COOKIE_SECURE=False,  # HTTPS 时改为 True\n'
        ')\n\n'
        'if __name__ == "__main__":\n'
        '    debug_mode = os.environ.get("FLASK_DEBUG", "false").lower() == "true"\n'
        '    app.run(debug=debug_mode, host="0.0.0.0", port=5000)'
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)

    # ── 4.6 ──
    add_heading_styled(doc, "4.6  路由权限鉴权补全", level=2)

    add_para(doc, "修复方案：", bold=True, size=11)
    add_para(doc,
             "（1）新增 @login_required 装饰器，未登录用户访问受保护路由时重定向至登录页。\n"
             "（2）首页 / 路由采用优雅降级策略：已登录显示用户信息，未登录显示"
             "「请先登录」提示 + 跳转按钮，不强制重定向以保持用户体验。\n"
             "（3）登录成功后执行 session.clear() 重置会话，防止会话固定攻击（Session Fixation）。",
             first_indent=0.74)

    add_para(doc, "核心代码（app.py）：", bold=True, size=10, color=(47, 84, 150))
    code = (
        'from functools import wraps\n\n'
        'def login_required(f):\n'
        '    @wraps(f)\n'
        '    def decorated(*args, **kwargs):\n'
        '        if "username" not in session:\n'
        '            return redirect(url_for("login"))\n'
        '        return f(*args, **kwargs)\n'
        '    return decorated\n\n'
        '# 登录成功后重置会话\n'
        'session.clear()\n'
        'session.permanent = True\n'
        'session["username"] = username\n'
        'session["csrf_token"] = secrets.token_hex(32)'
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 五、修复后功能与安全性测试
    # ══════════════════════════════════════════════════

    add_heading_styled(doc, "五、修复后功能与安全性测试", level=1)
    add_separator(doc)

    add_heading_styled(doc, "5.1  功能测试", level=2)

    make_table(doc,
               ["测试用例", "预期结果", "实际结果", "结论"],
               [
                   ["管理员登录（admin / admin123）", "登录成功，跳转至首页", "登录成功，正常跳转", "通过"],
                   ["普通用户登录（alice / alice2025）", "登录成功，跳转至首页", "登录成功，正常跳转", "通过"],
                   ["错误密码登录", "显示统一错误提示", "显示「用户名或密码错误」", "通过"],
                   ["空用户名 / 空密码提交", "显示非空提示", "显示错误提示", "通过"],
                   ["登录后用户信息展示", "显示脱敏后的信息", "用户名/角色/邮箱/脱敏手机/脱敏余额", "通过"],
                   ["退出登录", "清除会话，跳转首页", "跳转首页，显示「请先登录」", "通过"],
                   ["未登录访问首页", "显示登录提示", "显示「请先登录」+ 链接", "通过"],
                   ["CSRF 令牌缺失请求", "请求被拦截（400）", "返回 400 错误", "通过"],
                   ["XSS 注入测试（<script> 等）", "标签被过滤或转义", "标签被移除，页面正常", "通过"],
               ],
               col_widths=[4.5, 3.5, 3.5, 1.5])

    add_heading_styled(doc, "5.2  安全测试结果", level=2)

    make_table(doc,
               ["检测项", "修复前状态", "修复后状态", "结论"],
               [
                   ["密码明文存储", "存在（全部明文）", "已消除（scrypt 哈希）", "通过"],
                   ["前端密码泄露", "页面回显密码明文", "已消除（password 字段不传入模板）", "通过"],
                   ["HTML 注释泄露账号", "login.html 首行注释泄露", "已删除全部调试注释", "通过"],
                   ["手机号泄露", "完整展示 13800138000", "脱敏 138****8000", "通过"],
                   ["余额泄露", "完整展示 99999", "脱敏 ¥99,999.00", "通过"],
                   ["XSS 注入", "无过滤", "输入过滤 + Jinja2 转义", "通过"],
                   ["CSRF 防护", "无令牌校验", "before_request 全局校验", "通过"],
                   ["会话超时", "无", "2 小时自动过期", "通过"],
                   ["Cookie 安全", "默认设置", "HttpOnly + SameSite=Lax", "通过"],
                   ["会话固定攻击", "可被利用", "登录后 session.clear()", "通过"],
                   ["调试信息泄露", "debug=True 暴露调用栈", "环境变量控制，默认关闭", "通过"],
                   ["密钥强度", "dev-key-2025（弱）", "环境变量 / 256 位随机密钥", "通过"],
               ],
               col_widths=[3.5, 3.5, 4.5, 1.5])

    add_para(doc, "测试结论：", bold=True, size=11, space_before=8)
    add_para(doc,
             "修复完成后，系统全部原有业务功能正常可用，同时不存在明文密码存储与展示漏洞、"
             "前端源码无任何账号密码泄露、成功拦截 XSS 恶意脚本注入、表单具备 CSRF 防护能力、"
             "敏感数据已完成脱敏处理、服务器配置安全合规无调试信息泄露、会话加密安全无法被伪造劫持。",
             first_indent=0.74)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 六、修复前后安全对比分析
    # ══════════════════════════════════════════════════

    add_heading_styled(doc, "六、修复前后安全对比分析", level=1)
    add_separator(doc)

    make_table(doc,
               ["安全维度", "修复前", "修复后"],
               [
                   ["密码存储", "明文存储（== 字符串比对）", "Werkzeug scrypt 哈希（check_password_hash）"],
                   ["密码前端展示", "页面回显密码明文", "不传递 password 字段"],
                   ["默认账号泄露", "HTML 注释硬编码 admin/admin123", "全部删除"],
                   ["手机号展示", "完整展示 13800138000", "脱敏 138****8000"],
                   ["余额展示", "完整数字 99999", "货币格式 ¥99,999.00"],
                   ["XSS 防护", "无输入过滤", "HTML 过滤 + Jinja2 autoescape"],
                   ["CSRF 防护", "无校验", "secrets Token + before_request 拦截"],
                   ["会话固定攻击", "可被利用", "登录后 session.clear() 重置"],
                   ["会话超时", "永久有效", "2 小时自动过期"],
                   ["Cookie HttpOnly", "未设置", "已启用"],
                   ["Cookie SameSite", "未设置", "Lax 模式"],
                   ["Secret Key", "固定 dev-key-2025", "环境变量 / 256 位随机密钥"],
                   ["Debug 模式", "强制开启", "环境变量控制，默认关闭"],
                   ["异常处理", "无", "400/404/500 全局错误处理器"],
                   ["登录鉴权", "无装饰器", "@login_required + 路由鉴权"],
               ],
               col_widths=[3, 5.5, 6])

    add_para(doc, "总体评价：", bold=True, size=11, space_before=8)
    add_para(doc,
             "修复前：系统属于「裸奔」漏洞系统，存在多项高危漏洞，可被直接渗透、窃取账号、"
             "伪造请求、注入恶意代码、通过调试器实现远程代码执行，完全不具备任何线上部署条件。",
             first_indent=0.74)
    add_para(doc,
             "修复后：系统实现密码加密、输入过滤、请求防护、权限校验、数据脱敏、安全配置"
             "六大安全机制，修复全部 7 项已知漏洞。安全覆盖从存储层、传输层、展示层到配置层"
             "的完整攻击面，满足基础 Web 应用安全开发规范。",
             first_indent=0.74)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 七、安全优化总结与后续改进方案
    # ══════════════════════════════════════════════════

    add_heading_styled(doc, "七、安全优化总结与后续改进方案", level=1)
    add_separator(doc)

    add_heading_styled(doc, "7.1  实验总结", level=2)

    add_para(doc,
             "本次系统原始漏洞均为 Web 开发中典型的不安全编码问题，主要成因包括：开发规范缺失、"
             "安全意识不足、未采用加密存储、缺少基础安全防护机制。通过本次逐项审计与修复，"
             "成功消除全部高危安全隐患，在保证业务完整性的前提下，显著提升了 Web 应用的抗攻击能力。",
             first_indent=0.74)

    add_para(doc, "核心收获：", bold=True, size=11)
    add_bullet(doc, "密码必须在存储层开始加密，任何明文存储都是不可接受的安全红线")
    add_bullet(doc, "前端展示的数据必须遵循「最小必要」原则，敏感字段必须脱敏")
    add_bullet(doc, "所有 POST 请求必须配置 CSRF 防护，这是表单应用的标配")
    add_bullet(doc, "用户输入永远不可信，输入过滤 + 输出转义是 XSS 防护的双保险")
    add_bullet(doc, "开发密钥、调试模式等配置应通过环境变量管理，杜绝硬编码")
    add_bullet(doc, "会话管理需要超时机制 + 安全 Cookie 标志 + 防固定攻击")

    add_heading_styled(doc, "7.2  后续优化方向", level=2)

    add_para(doc,
             "虽然本项目已完成基础安全加固，但在生产环境中还应考虑以下升级方向：",
             first_indent=0.74)

    make_table(doc,
               ["优先级", "优化项", "具体方案"],
               [
                   ["P0", "数据库化存储", "将 USERS 字典替换为 SQLite / PostgreSQL + SQLAlchemy ORM，实现数据持久化与事务管理"],
                   ["P0", "HTTPS 强制", "配置 Nginx 反向代理 + Let's Encrypt 证书，同时启用 SESSION_COOKIE_SECURE=True"],
                   ["P1", "登录频率限制", "引入 Flask-Limiter，限制单 IP 每分钟登录尝试次数，抵御暴力破解攻击"],
                   ["P1", "验证码接入", "登录页增加图形验证码或 Google reCAPTCHA，区分人为操作与自动化脚本"],
                   ["P1", "日志审计系统", "记录登录成功/失败日志，对接集中式日志平台（如 ELK），支持异常行为告警"],
                   ["P2", "Flask-Login 集成", "使用官方扩展代替手写 login_required 装饰器，获得更完善的会话管理能力"],
                   ["P2", "Flask-WTF 表单", "使用 WTForms 统一管理表单验证与 CSRF，减少手写校验代码"],
                   ["P2", "内容安全策略", "设置 Content-Security-Policy 响应头，限制可加载的资源来源"],
                   ["P3", "双因素认证", "增加 TOTP（基于时间的一次性密码）或短信验证码二次验证"],
                   ["P3", "密码复杂度策略", "增加密码长度/字符类型校验、定期更换提醒、历史密码防重用"],
               ],
               col_widths=[1.5, 3, 10])

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 八、实验心得
    # ══════════════════════════════════════════════════

    add_heading_styled(doc, "八、实验心得", level=1)
    add_separator(doc)

    add_para(doc,
             "通过本次 Flask 用户管理系统漏洞修复实验，我系统掌握了 Web 常见漏洞的产生原理、"
             "漏洞复现方式与标准化修复方案。",
             first_indent=0.74)

    add_para(doc,
             "在代码审计阶段，通过对原始代码逐行分析，我发现看似「简洁」的代码背后往往隐藏着"
             "大量的安全风险。明文存储密码仅是为了减少几行代码，却导致了最严重的数据泄露漏洞；"
             "调试注释仅是为了开发方便，却直接向所有用户暴露了管理员账号。这些案例警示我们："
             "安全不应该在开发结束后才考虑，而是必须在每一行代码的编写过程中贯彻。",
             first_indent=0.74)

    add_para(doc,
             "在漏洞修复阶段，我深入理解了多项安全机制的工作原理：Werkzeug 的密码哈希函数"
             "如何通过加盐和迭代计算防止彩虹表攻击；CSRF 令牌如何利用同源策略和不可预测性"
             "防御跨站请求；输入过滤与输出转义如何协同构建 XSS 的多层防线。这些理解远超"
             "简单的 API 调用层面，而是深入到了安全设计的底层逻辑。",
             first_indent=0.74)

    add_para(doc,
             "在回归测试阶段，我体会到安全加固必须建立在充分的功能测试基础上。每一处安全修改"
             "都可能影响正常的业务流程，必须在修复后逐条验证原有功能是否正常。安全与业务"
             "不是对立关系，高质量的安全加固应当在「看不见」的地方保护系统，而不打扰用户。",
             first_indent=0.74)

    add_para(doc,
             "深刻认识到明文存储、输入过滤缺失、配置不规范等低级编码问题会造成极大的安全风险。"
             "在今后的 Web 开发中，将严格遵循安全开发规范，优先考虑数据加密、输入校验、权限控制"
             "与请求防护，从代码底层规避安全漏洞，提升 Web 项目的安全性与健壮性。",
             first_indent=0.74)

    # ── 文档结尾 ──
    doc.add_paragraph()
    add_separator(doc)

    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run("— 报告完 —")
    set_run_font(run, size=11, color=(150, 150, 150))

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "本报告仅供安全教学与技术交流使用。报告中涉及的漏洞代码已全部修复，\n"
        "请勿将默认测试账号部署到公开生产环境。"
    )
    set_run_font(run, size=9, color=(180, 180, 180))

    # ── 保存 ──
    doc.save(output_path)
    print(f"报告已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "Flask用户信息管理平台_漏洞检测与安全修复报告.docx")
    generate_report(out)
