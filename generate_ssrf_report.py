#!/usr/bin/env python3
"""
URL 抓取功能 — SSRF 服务端请求伪造漏洞复现及修复报告
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="微软雅黑", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return heading


def add_para(doc, text, bold=False, size=11, color=None, align=None,
             space_after=6, space_before=0, first_indent=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = Pt(20)
    if align:
        para.alignment = align
    if first_indent:
        para.paragraph_format.first_line_indent = Cm(first_indent)
    return para


def add_code_block(doc, code, indent=1.0):
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)


def add_separator(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def make_table(doc, headers, rows, col_widths=None, header_color="C0392B"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def generate_report(output_path):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(11)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ═══════════ 封面 ═══════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Flask 用户信息管理平台")
    set_run_font(run, size=26, bold=True, color=(192, 57, 43))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("URL 抓取功能 —— SSRF 服务端请求伪造漏洞\n复现及修复报告")
    set_run_font(run, size=18, bold=False, color=(89, 89, 89))

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("-" * 40)
    set_run_font(run, size=12, color=(200, 200, 200))
    doc.add_paragraph()

    for text in [
        "实验性质：SSRF 服务端请求伪造漏洞复现与修复实训",
        "目标系统：Flask 用户信息管理系统（/fetch-url 路由）",
        "文档版本：V1.0 终审版",
        "生成日期：2026 年 7 月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=12, color=(100, 100, 100))

    doc.add_page_break()

    # ═══════════ 目录 ═══════════
    add_heading_styled(doc, "目  录", level=1)
    doc.add_paragraph()
    toc_items = [
        "一、实验概述",
        "二、漏洞总览",
        "三、漏洞 1：file:// 协议读取服务器本地文件",
        "    3.1  复现内容",
        "    3.2  修复方案",
        "四、漏洞 2：访问 127.0.0.1 本地内网端口",
        "    4.1  复现内容",
        "    4.2  修复方案",
        "五、漏洞 3：扫描内网网段资产测绘",
        "    5.1  复现内容",
        "    5.2  修复方案",
        "六、修复方案汇总与安全对比",
        "七、修复后安全验证",
        "八、实验总结与心得",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(item)
        set_run_font(run, size=11, bold=not item.startswith("    "))

    doc.add_page_break()

    # ═══════════ 一、实验概述 ═══════════
    add_heading_styled(doc, "一、实验概述", level=1)
    add_separator(doc)

    add_para(doc,
             "本次实验针对 Flask 用户信息管理平台中的 URL 抓取功能 /fetch-url 进行安全审计。"
             "该路由使用 urllib.request.urlopen() 直接访问用户提交的 URL，未对协议、目标 IP、"
             "端口做任何限制，导致存在严重的高危 SSRF 服务端请求伪造漏洞。",
             first_indent=0.74)

    add_para(doc,
             "本次实验共发现 3 项高危子漏洞：file:// 协议读取本地文件、127.0.0.1 内网端口探测、"
             "内网网段资产测绘。攻击者无需高权限即可通过服务器发起内网请求，危害极大。",
             first_indent=0.74)

    add_para(doc,
             "实验流程包括：漏洞复现（浏览器手工操作）-> 根因分析 -> 安全加固 -> "
             "修复后验证，完整覆盖从攻击到防御的全生命周期。",
             first_indent=0.74)

    # ═══════════ 二、漏洞总览 ═══════════
    add_heading_styled(doc, "二、漏洞总览", level=1)
    add_separator(doc)

    make_table(doc,
               ["编号", "漏洞名称", "漏洞类型", "危害等级"],
               [
                   ["VUL-SSRF-01", "file:// 协议读取服务器本地文件", "SSRF 任意文件读取", "高危"],
                   ["VUL-SSRF-02", "访问 127.0.0.1 本地内网端口", "SSRF 内网探测", "高危"],
                   ["VUL-SSRF-03", "扫描内网网段资产测绘", "SSRF 内网穿透", "高危"],
               ],
               col_widths=[2.5, 5.5, 3.5, 1.5])

    add_para(doc, "漏洞根因总结：", bold=True, size=11, space_before=6)
    add_para(doc, "（1）未限制 URL 协议，file:// 可直接读取本地任意文件；", first_indent=0.74)
    add_para(doc, "（2）未解析目标 IP，127.0.0.1 等内网地址可被直接请求；", first_indent=0.74)
    add_para(doc, "（3）未限制访问端口，22/3306/6379 等高危服务可被探测；", first_indent=0.74)
    add_para(doc, "（4）未限制内网网段，192.168/10/172.16 均可扫描。", first_indent=0.74)

    add_para(doc, "漏洞代码片段（修复前）：", bold=True, size=10, color=(192, 57, 43))
    add_code_block(doc,
        '# 修复前：无协议/IP/端口校验，直接 urlopen\n'
        '@app.route("/fetch-url", methods=["POST"])\n'
        "def fetch_url():\n"
        '    target_url = request.form.get("url", "")\n'
        "    # 无任何校验\n"
        "    resp = urllib.request.urlopen(target_url)  # file:// 也可执行"
    )

    # ═══════════════════════════════════════════
    # 漏洞 1
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "三、漏洞 1：file:// 协议读取服务器本地文件", level=1)
    add_separator(doc)

    add_heading_styled(doc, "3.1  复现内容", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "使用任意账号登录网站，进入首页。",
        "找到页面 URL 抓取输入框，填入 Payload（区分系统）。",
        "Linux 服务器：file:///app.py",
        "Windows 服务器：file://C:/app.py",
        '点击「抓取」按钮提交 POST 请求。',
        "查看页面下方抓取结果区域。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "服务器成功读取本地 app.py 源代码，页面完整输出后端源码内容。"
             "更换 file:///etc/passwd（Linux）/ file://C:/Windows/System32/drivers/etc/hosts（Windows）"
             "可读取系统敏感文件，SSRF 文件读取漏洞存在。",
             first_indent=0.74)

    add_para(doc, "攻击原理示意图：", bold=True, size=11)
    add_code_block(doc,
        "用户提交: file:///app.py\n"
        "urllib.urlopen(\"file:///app.py\")\n"
        "-> 打开本地文件 /app.py\n"
        "-> 页面输出：app.py 全部源代码\n\n"
        "用户提交: file:///etc/passwd\n"
        "-> 打开系统用户文件\n"
        "-> 页面输出：root:x:0:0:root:/root:/bin/bash"
    )

    add_heading_styled(doc, "3.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "使用 urlparse 解析 URL 协议，设置协议白名单仅允许 http 和 https，"
             "拦截 file:// 等所有非 http 协议。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "urllib.parse.urlparse() 可准确提取 URL 的 scheme 部分。"
             "通过白名单校验，仅放行 scheme 为 http 或 https 的 URL。"
             "file:// 协议因不在白名单中，直接被拒绝，不会进入 urlopen 执行阶段。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        '# 协议白名单\n'
        'ALLOWED_URL_SCHEMES = ("http", "https")\n\n'
        'parsed = urllib.parse.urlparse(url_str)\n'
        "if not parsed.scheme or not parsed.netloc:\n"
        '    return False, "URL 格式不合法"\n'
        "if parsed.scheme not in ALLOWED_URL_SCHEMES:\n"
        '    return False, "仅支持 http 和 https 协议"'
    )

    # ═══════════════════════════════════════════
    # 漏洞 2
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "四、漏洞 2：访问 127.0.0.1 本地内网端口", level=1)
    add_separator(doc)

    add_heading_styled(doc, "4.1  复现内容", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "登录账号进入首页 URL 抓取功能。",
        "输入 Payload：http://127.0.0.1:3306（MySQL 端口），提交抓取。",
        "更换端口 http://127.0.0.1:6379（Redis）、http://127.0.0.1:8080（本地后台）分别测试。",
        "观察页面返回状态码、响应内容、超时差异。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "服务器主动向本机内网端口发起请求，根据返回内容 / 超时可判断端口是否开放，"
             "攻击者可批量扫描本机全部端口，本地内网探测 SSRF 漏洞存在。",
             first_indent=0.74)

    add_para(doc, "可探测的内网服务：", bold=True, size=11)
    make_table(doc,
               ["端口", "常见服务", "攻击利用"],
               [
                   ["22", "SSH", "暴力破解内网 SSH"],
                   ["80/443", "Web 后台", "访问隐藏管理后台"],
                   ["3306", "MySQL", "探测数据库版本/弱口令"],
                   ["6379", "Redis", "未授权访问/写 crontab"],
                   ["8080", "代理/后台", "访问内网 Web 管理界面"],
                   ["27017", "MongoDB", "未授权数据读取"],
               ],
               col_widths=[2.5, 4, 6.5])

    add_heading_styled(doc, "4.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "DNS 解析目标域名获取真实 IP，黑名单拦截全部内网私有网段。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "通过 socket.gethostbyname() 将域名解析为 IP 地址，然后使用 ipaddress 模块"
             "判断 IP 是否属于内网保留网段。127.0.0.1 解析后属于 127.0.0.0/8 网段，"
             "直接被拦截。内网 IP 不会进入 urlopen 执行阶段。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        '# 内网 IP 网段黑名单\n'
        'INTERNAL_IP_NETWORKS = [\n'
        '    "127.0.0.0/8",     # 本机回环\n'
        '    "10.0.0.0/8",      # A 类私有\n'
        '    "172.16.0.0/12",   # B 类私有\n'
        '    "192.168.0.0/16",  # C 类私有\n'
        '    "169.254.0.0/16",  # 链路本地\n'
        "]\n\n"
        "ip = socket.gethostbyname(hostname)\n"
        "ip_obj = ipaddress.ip_address(ip)\n"
        "for network_str in INTERNAL_IP_NETWORKS:\n"
        "    if ip_obj in ipaddress.ip_network(network_str, strict=False):\n"
        '        return False, "不允许访问内网地址"'
    )

    # ═══════════════════════════════════════════
    # 漏洞 3
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "五、漏洞 3：扫描内网网段资产测绘", level=1)
    add_separator(doc)

    add_heading_styled(doc, "5.1  复现内容", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "保持登录状态，在抓取框依次输入内网 Payload。",
        "http://192.168.1.1:80、http://192.168.1.10:80、http://10.0.0.5:8080。",
        "逐个提交抓取请求，对比页面返回内容与响应时长。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "服务器可访问内网所有私有网段主机，通过响应差异判断内网存活设备、开放 Web 端口，"
             "可对内网进行完整资产测绘，高危内网穿透 SSRF 漏洞存在。",
             first_indent=0.74)

    add_para(doc, "内网穿透攻击链路：", bold=True, size=11)
    add_code_block(doc,
        "外网攻击者 -> 提交 http://192.168.1.10:80\n"
        "-> 服务器发起请求到内网 192.168.1.10\n"
        "-> 获取内网 Web 页面内容\n"
        "-> 攻击者间接获取内网资产信息\n\n"
        "批量扫描:\n"
        "http://192.168.1.1:3306  -> MySQL 端口开放?  -> 响应差异\n"
        "http://10.0.0.5:6379    -> Redis 开放?       -> 响应差异\n"
        "http://172.16.0.1:22    -> SSH 开放?          -> 超时/拒绝"
    )

    add_heading_styled(doc, "5.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "三重校验联动防御：协议白名单 + 端口白名单 + 内网 IP 拦截，组合后完全阻断 SSRF 攻击链。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "端口白名单仅允许 80 和 443，阻断非 Web 端口的访问尝试。"
             "即使攻击者使用 http://192.168.1.10:6379，也会因 6379 不在端口白名单中被拦截。"
             "三重校验叠加后，攻击者无法通过任何协议、端口、网段发起 SSRF 攻击。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        '# 端口白名单\n'
        'ALLOWED_PORTS = {80, 443}\n\n'
        "if port not in ALLOWED_PORTS:\n"
        '    return False, "仅允许访问标准 Web 端口（80、443）"\n\n'
        "# 完整校验流程\n"
        "def validate_fetch_url(url_str):\n"
        "    # 1. URL 格式校验\n"
        "    # 2. 协议白名单（仅 http/https）\n"
        "    # 3. 端口白名单（仅 80/443）\n"
        "    # 4. DNS 解析 + 内网 IP 拦截\n"
        "    return True, None  # 全部通过"
    )

    # ═══════════════════════════════════════════
    # 六、修复汇总
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "六、修复方案汇总与安全对比", level=1)
    add_separator(doc)

    add_para(doc, "三重防御体系：", bold=True, size=11, space_before=6)
    make_table(doc,
               ["防御层", "校验内容", "拦截目标"],
               [
                   ["第 1 层：协议白名单", "scheme in (http, https)", "file://, dict://, ftp://, gopher://"],
                   ["第 2 层：端口白名单", "port in (80, 443)", "22, 3306, 6379, 8080 等"],
                   ["第 3 层：内网 IP 拦截", "IP not in 内网段", "127.0.0.1, 192.168.x.x, 10.x.x.x"],
               ],
               col_widths=[3, 4, 6])

    make_table(doc,
               ["安全维度", "修复前", "修复后"],
               [
                   ["协议限制", "无限制，支持 file:// 等", "白名单仅 http/https"],
                   ["端口限制", "无限制，任意端口", "白名单仅 80/443"],
                   ["内网 IP 检测", "不解析 IP，直接请求", "DNS 解析 + 5 个内网段拦截"],
                   ["本地文件读取", "file:// 可读取任意文件", "协议白名单拦截 file://"],
                   ["内网端口探测", "127.0.0.1:3306 可探测", "端口 + IP 双重拦截"],
                   ["内网资产扫描", "192.168/10 全段可扫", "内网 IP 网段拦截"],
                   ["错误信息", "暴露内网连接详情", "统一错误提示"],
                   ["CSRF 防护", "排除在 CSRF 校验外", "纳入 CSRF Token 校验"],
               ],
               col_widths=[3, 4.5, 5.5])

    # ═══════════════════════════════════════════
    # 七、修复后验证
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "七、修复后安全验证", level=1)
    add_separator(doc)

    make_table(doc,
               ["测试用例", "预期结果", "实际结果", "结论"],
               [
                   ["正常 http://example.com", "返回页面内容", "200, Example Domain", "通过"],
                   ["正常 https://example.com", "返回页面内容", "正常返回", "通过"],
                   ["file:///etc/passwd", "协议不合法，拦截", "仅支持 http/https", "通过"],
                   ["file:///app.py", "协议不合法，拦截", "URL 格式不合法", "通过"],
                   ["dict://127.0.0.1:6379", "协议不合法，拦截", "仅支持 http/https", "通过"],
                   ["ftp://example.com", "协议不合法，拦截", "仅支持 http/https", "通过"],
                   ["http://127.0.0.1:80", "内网地址拦截", "不允许访问内网地址", "通过"],
                   ["http://192.168.1.1:80", "内网地址拦截", "不允许访问内网地址", "通过"],
                   ["http://10.0.0.1:80", "内网地址拦截", "不允许访问内网地址", "通过"],
                   ["http://example.com:22", "非标准端口拦截", "仅允许 80/443", "通过"],
                   ["http://example.com:3306", "非标准端口拦截", "仅允许 80/443", "通过"],
                   ["http://example.com:6379", "非标准端口拦截", "仅允许 80/443", "通过"],
                   ["空 URL", "格式不合法拦截", "URL 不能为空", "通过"],
                   ["CSRF Token 缺失", "400 拦截", "400 错误", "通过"],
                   ["CSRF Token 错误", "400 拦截", "400 错误", "通过"],
               ],
               col_widths=[4, 3.5, 3.5, 1.5])

    add_para(doc, "验证结论：", bold=True, size=11, space_before=8)
    add_para(doc,
             "全部 3 项 SSRF 子漏洞已成功修复。正常 URL 抓取功能完全保留，"
             "所有攻击向量全部被拦截。修复后系统满足以下安全要求：\n"
             "（1）file:// 等危险协议被白名单拦截，无法读取本地文件；\n"
             "（2）127.0.0.1 等内网地址经 DNS 解析后被网段拦截，无法访问；\n"
             "（3）22/3306/6379 等高危端口被端口白名单拦截，无法探测；\n"
             "（4）192.168/10/172.16 等内网网段全部拦截，无法资产测绘；\n"
             "（5）错误信息统一化，不暴露内网响应细节。",
             first_indent=0.74)

    # ═══════════════════════════════════════════
    # 八、总结
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "八、实验总结与心得", level=1)
    add_separator(doc)

    add_para(doc,
             "通过本次 SSRF 漏洞复现与修复实验，深入理解了服务端请求伪造攻击的多种利用手法"
             "及其对应的分层防御策略。",
             first_indent=0.74)

    add_para(doc, "核心认知：", bold=True, size=11)
    add_para(doc,
             "（1）SSRF 的本质是「服务端发起了用户可控的网络请求」。当应用需要根据用户输入"
             "的 URL 去请求外部资源时，如果不对目标协议、IP、端口做限制，服务器就变成了攻击者的《跳板机》。",
             first_indent=0.74)

    add_para(doc,
             "（2）协议白名单是防御 SSRF 的第一道防线。file:// 协议可以直接读取服务器文件，"
             "dict:///gopher:// 等协议可以攻击内网 Redis 等服务。仅允许 http/https 可阻断这类攻击。",
             first_indent=0.74)

    add_para(doc,
             "（3）DNS 解析 + 内网 IP 核验是防御 SSRF 的关键技术。仅校验域名前缀（如检查"
             "是否包含 127.0.0.1）可被域名解析绕过。必须通过 socket.gethostbyname 获取真实 IP，"
             "再使用 ipaddress 模块做网段归属判断。",
             first_indent=0.74)

    add_para(doc,
             "（4）端口白名单可限制 SSRF 的攻击面。即使攻击者控制了内外网可达的地址，"
             "端口限制使其无法攻击非 Web 服务（MySQL 3306、Redis 6379 等），大大降低危害。",
             first_indent=0.74)

    add_para(doc,
             "通过本次实验深刻认识到，任何由服务端发起网络请求的功能都应被视为高风险功能。"
             "协议、IP、端口三重校验缺一不可，只有分层防御才能有效抵御 SSRF 攻击。",
             first_indent=0.74)

    # 结尾
    doc.add_paragraph()
    add_separator(doc)

    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run("报告完")
    set_run_font(run, size=11, color=(150, 150, 150))

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "本报告仅供安全教学与技术交流使用。报告中涉及的漏洞代码已全部修复，\n"
        "请勿将存在漏洞的代码版本部署到任何生产环境。"
    )
    set_run_font(run, size=9, color=(180, 180, 180))

    doc.save(output_path)
    print(f"报告已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "URL抓取功能_SSRF服务端请求伪造漏洞复现及修复报告.docx")
    generate_report(out)
