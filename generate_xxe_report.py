#!/usr/bin/env python3
"""
XML 数据导入功能 — XXE 外部实体注入漏洞复现及修复报告
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="微软雅黑", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return heading


def add_para(doc, text, bold=False, size=11, color=None, align=None,
             space_after=6, space_before=0, first_indent=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = Pt(20)
    if align:
        para.alignment = align
    if first_indent:
        para.paragraph_format.first_line_indent = Cm(first_indent)
    return para


def add_code_block(doc, code, indent=1.0):
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)


def add_separator(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def make_table(doc, headers, rows, col_widths=None, header_color="6C3483"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def generate_report(output_path):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(11)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 封面
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Flask 用户信息管理平台")
    set_run_font(run, size=26, bold=True, color=(108, 52, 131))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("XML 数据导入功能\nXXE 外部实体注入漏洞\n复现及修复报告")
    set_run_font(run, size=18, bold=False, color=(89, 89, 89))

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("-" * 40)
    set_run_font(run, size=12, color=(200, 200, 200))
    doc.add_paragraph()

    for text in [
        "实验性质：XXE 外部实体注入漏洞复现与修复实训",
        "目标系统：Flask 用户信息管理系统（/xml-import 路由）",
        "文档版本：V1.0 终审版",
        "生成日期：2026 年 7 月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=12, color=(100, 100, 100))

    doc.add_page_break()

    # 目录
    add_heading_styled(doc, "目  录", level=1)
    doc.add_paragraph()
    toc_items = [
        "一、实验概述",
        "二、漏洞总览",
        "三、漏洞详情与复现",
        "    3.1  复现 1：XXE 读取项目后端源码 app.py",
        "    3.2  复现 2：Linux 读取 /etc/passwd",
        "    3.3  复现 3：Windows 读取 hosts 文件",
        "    3.4  复现 4：多层路径遍历读取配置文件",
        "四、漏洞危害总结",
        "五、漏洞根因分析",
        "六、修复方案",
        "七、修复汇总与安全对比",
        "八、修复后安全验证",
        "九、实验总结与心得",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(item)
        set_run_font(run, size=11, bold=not item.startswith("    "))

    doc.add_page_break()

    # 一、实验概述
    add_heading_styled(doc, "一、实验概述", level=1)
    add_separator(doc)

    add_para(doc,
             "本次实验针对 Flask 用户信息管理平台中的 XML 数据导入功能 /xml-import 进行安全审计。"
             "该路由接收用户提交的 XML 文本并解析，但未禁用 XML 外部实体解析，"
             "导致存在严重的 XXE 外部实体注入漏洞（CWE-611），可被利用读取服务器任意本地文件。",
             first_indent=0.74)

    add_para(doc,
             "本次实验共发现 4 种 XXE 利用方式：读取项目源码、读取系统账号文件、读取 Windows hosts 文件、"
             "多层路径遍历读取配置文件。攻击者登录后即可读取服务器任意文件。",
             first_indent=0.74)

    # 二、漏洞总览
    add_heading_styled(doc, "二、漏洞总览", level=1)
    add_separator(doc)

    make_table(doc,
               ["编号", "攻击方式", "Payload 示例", "攻击效果"],
               [
                   ["VUL-XXE-01", "读取项目源码",
                    '<!ENTITY file SYSTEM "../app.py">', "泄露 app.py 全部后端代码"],
                   ["VUL-XXE-02", "Linux 系统文件",
                    '<!ENTITY passwd SYSTEM "/etc/passwd">', "读取系统账号、UID、Shell"],
                   ["VUL-XXE-03", "Windows 系统文件",
                    '<!ENTITY hosts SYSTEM "C:/Windows/.../hosts">', "获取内网域名映射"],
                   ["VUL-XXE-04", "多层路径遍历",
                    '<!ENTITY config SYSTEM "../../config">', "读取密钥和数据库配置"],
               ],
               col_widths=[2, 3, 5, 3.5])

    add_para(doc, "漏洞根因一句话总结：", bold=True, size=11, color=(108, 52, 131), space_before=6)
    add_para(doc,
             "后端未禁用 XML 外部实体解析 + 自动提取 SYSTEM 路径并读取本地文件 + 读取内容直接回显 = "
             "XXE 任意文件读取漏洞。",
             first_indent=0.74)

    add_para(doc, "漏洞代码片段（修复前）：", bold=True, size=10, color=(192, 57, 43))
    add_code_block(doc,
        '# 修复前：三错叠加导致 XXE\n'
        'xml_data = request.form.get("xml_data", "")\n\n'
        "# 错误 1: 未禁用外部实体解析\n"
        "root = ET.fromstring(xml_data)  # 默认允许外部实体\n\n"
        "# 错误 2: 自动读取 SYSTEM 文件\n"
        'if "<!ENTITY" in xml_data and "SYSTEM" in xml_data:\n'
        '    file_path = re.search(...).group(1)  # 提取路径\n'
        "    with open(file_path, \"r\") as f:     # 读取本地文件\n"
        "        file_content = f.read()\n"
        '        xml_data = xml_data.replace("&xxe;", file_content)\n\n'
        "# 错误 3: 文件内容嵌入解析结果返回前端\n"
        "result_json = json.dumps(users)  # 文件内容出现在 JSON 中"
    )

    # 三、漏洞详情与复现
    doc.add_page_break()
    add_heading_styled(doc, "三、漏洞详情与复现", level=1)
    add_separator(doc)

    add_para(doc, "环境前提：", bold=True, size=11)
    add_para(doc,
             "靶场全部功能正常运行，登录后导航栏存在「XML 导入」入口，/xml-import 接口无 XXE 防护，"
             "自动读取 SYSTEM 指向的本地文件并回显。",
             first_indent=0.74)

    # 3.1
    add_heading_styled(doc, "3.1  复现 1：XXE 读取项目后端源码 app.py", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "登录任意账号，点击导航栏「XML 导入」进入 xml_import.html 页面。",
        "在 XML 文本框输入恶意 XXE 载荷（含 <!ENTITY file SYSTEM \"../app.py\">）。",
        "点击「导入」按钮提交 POST 请求。",
        "查看下方 pre 标签内 JSON 解析结果。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "恶意 XML 载荷：", bold=True, size=10, color=(192, 57, 43))
    add_code_block(doc,
        '<?xml version="1.0"?>\n'
        '<!DOCTYPE user [\n'
        '<!ENTITY file SYSTEM "../app.py">\n'
        "]>\n"
        "<user>\n"
        "  <name>&file;</name>\n"
        "  <email>test@test.com</email>\n"
        "</user>"
    )

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "后端识别 <!ENTITY SYSTEM，读取 ../app.py 源码内容，将完整代码替换 &file; 实体，"
             "最终 JSON 中 name 字段输出 app.py 全部源代码，源码泄露，XXE 漏洞可利用。",
             first_indent=0.74)

    add_para(doc, "攻击原理示意图：", bold=True, size=11)
    add_code_block(doc,
        "1. 用户提交含 ENTITY 的 XML\n"
        "2. 后端检测到 SYSTEM \"../app.py\"\n"
        "3. open(\"../app.py\") 读取源码\n"
        "4. &file; 替换为 app.py 全部内容\n"
        "5. JSON 输出: {\"name\": \"from flask ...\", \"email\": \"...\"}"
    )

    # 3.2
    add_heading_styled(doc, "3.2  复现 2：Linux 环境读取系统账号文件 /etc/passwd", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "进入 XML 导入页面，输入含 /etc/passwd 的 XXE 载荷。",
        "提交导入，查看 JSON 输出。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "恶意 XML 载荷：", bold=True, size=10, color=(192, 57, 43))
    add_code_block(doc,
        '<?xml version="1.0"?>\n'
        '<!DOCTYPE user [\n'
        '<!ENTITY passwd SYSTEM "/etc/passwd">\n'
        "]>\n"
        "<user>\n"
        "  <name>test</name>\n"
        "  <email>&passwd;</email>\n"
        "</user>"
    )

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "服务器读取 /etc/passwd 系统用户文件，所有系统账号、权限信息完整展示在返回 JSON 内，"
             "系统敏感信息泄露。",
             first_indent=0.74)

    # 3.3
    add_heading_styled(doc, "3.3  复现 3：Windows 环境读取 hosts 文件", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "XML 输入框填写 Windows hosts 文件 XXE 载荷。",
        "提交表单查看解析结果。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "恶意 XML 载荷：", bold=True, size=10, color=(192, 57, 43))
    add_code_block(doc,
        '<?xml version="1.0"?>\n'
        '<!DOCTYPE user [\n'
        '<!ENTITY hosts SYSTEM "C:/Windows/System32/drivers/etc/hosts">\n'
        "]>\n"
        "<user>\n"
        "  <name>&hosts;</name>\n"
        "  <email>demo@test.com</email>\n"
        "</user>"
    )

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "成功读取本地 hosts 文件，内网 IP、域名映射信息全部回显。",
             first_indent=0.74)

    # 3.4
    add_heading_styled(doc, "3.4  复现 4：多层路径遍历读取配置文件", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "输入带多层 ../ 穿越载荷。",
        "提交导入。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "路径穿越生效，读取项目根目录无后缀配置文件，数据库密钥、后台配置全部暴露。",
             first_indent=0.74)

    # 四、危害总结
    doc.add_page_break()
    add_heading_styled(doc, "四、漏洞危害总结", level=1)
    add_separator(doc)

    add_para(doc,
             "攻击者仅需登录账号，构造携带外部实体的 XML 文本，即可通过 XXE 漏洞读取服务器任意本地文件，"
             "泄露后端源码、数据库密钥、系统账号、内网配置等核心敏感数据，属于高危信息泄露漏洞。",
             first_indent=0.74)

    make_table(doc,
               ["攻击目标", "Payload SYSTEM 路径", "泄露内容"],
               [
                   ["项目源码", "../app.py", "Flask 路由、密钥、数据库配置"],
                   ["Linux 系统账号", "/etc/passwd", "全部系统用户、UID、GID、Shell"],
                   ["Linux Shadow", "/etc/shadow", "用户密码哈希（提权用）"],
                   ["数据库配置", "../config", "MySQL/Redis 连接凭据"],
                   ["SSH 密钥", "~/.ssh/id_rsa", "服务器 SSH 私钥"],
                   ["Windows hosts", "C:/Windows/.../hosts", "内网域名解析映射"],
                   ["项目数据库", "data/users.db", "SQLite 全部用户数据"],
               ],
               col_widths=[3, 4, 5])

    add_para(doc, "CVSS 3.1 评分：9.1（Critical）", bold=True, size=12, color=(108, 52, 131), space_before=6)

    # 五、根因分析
    doc.add_page_break()
    add_heading_styled(doc, "五、漏洞根因分析", level=1)
    add_separator(doc)

    add_para(doc, "三个致命错误的叠加导致 XXE：", bold=True, size=11)

    add_para(doc, "错误 1：未禁用 XML 外部实体解析", bold=True, size=11, color=(192, 57, 43))
    add_para(doc,
             "Python 的 xml.etree.ElementTree.fromstring() 在解析 XML 时，如果 XML 文本中包含 "
             "<!DOCTYPE> 和 <!ENTITY ... SYSTEM ...> 定义，默认会尝试解析外部实体。虽然 Python 3.x "
             "默认不解析外部实体加载，但本靶场通过正则提取 SYSTEM 路径后主动 open() 读取文件，"
             "绕过了 Python 自身的默认保护。",
             first_indent=0.74)

    add_para(doc, "错误 2：主动提取文件路径并读取本地文件", bold=True, size=11, color=(192, 57, 43))
    add_para(doc,
             "通过正则 r'<!ENTITY\\s+\\w+\\s+SYSTEM\\s+\"([^\"]+)\"' 从 XML 中提取文件路径，"
             "然后使用 open() 直接读取。这实际上重新实现了 XXE 的文件读取功能，"
             "完全绕过了 Python 的 XML 解析器安全限制。",
             first_indent=0.74)

    add_para(doc, "错误 3：读取内容直接嵌入解析结果返回前端", bold=True, size=11, color=(192, 57, 43))
    add_para(doc,
             "读取到的文件内容通过字符串替换注入到 XML 文本中，然后解析为 JSON 返回给用户。"
             "攻击者无需任何额外操作即可在前端看到完整的文件内容。",
             first_indent=0.74)

    # 六、修复方案
    doc.add_page_break()
    add_heading_styled(doc, "六、修复方案", level=1)
    add_separator(doc)

    add_para(doc, "修复 1：拦截含外部实体定义的恶意 XML", bold=True, size=11, color=(39, 174, 96))
    add_para(doc,
             "对 XML 文本进行安全预检，检查是否包含 <!ENTITY、SYSTEM、PUBLIC 等危险关键字。"
             "如果检测到这些关键字，直接拒绝解析并返回明确错误提示。",
             first_indent=0.74)

    add_para(doc, "修复后代码：", bold=True, size=10, color=(39, 174, 96))
    add_code_block(doc,
        '# 修复后：拦截含外部实体的恶意 XML\n'
        'if "<!ENTITY" in xml_data.upper():\n'
        '    error = "XML 中包含不安全的实体定义，已拒绝解析"\n\n'
        'elif "SYSTEM" in xml_data or "PUBLIC" in xml_data:\n'
        '    error = "XML 中包含不安全的实体引用，已拒绝解析"\n\n'
        "else:\n"
        "    # 仅解析纯净 XML 结构\n"
        "    root = ET.fromstring(xml_data)"
    )

    add_para(doc, "修复 2：删除自动读取本地文件的高危逻辑", bold=True, size=11, color=(39, 174, 96))
    add_para(doc,
             "完全删除正则提取 SYSTEM 路径、open() 读取文件、字符串替换实体引用的全部代码。"
             "XML 解析仅处理用户提交的 XML 文本中已有的数据，不访问任何外部资源。",
             first_indent=0.74)

    add_para(doc, "修复 3：统一错误信息", bold=True, size=11, color=(39, 174, 96))
    add_para(doc,
             "XML 解析错误的提示从暴露具体错误信息（如行号、列号）改为统一提示"
             "「XML 格式错误，请检查后重试」，防止信息泄露。",
             first_indent=0.74)

    # 七、修复汇总
    doc.add_page_break()
    add_heading_styled(doc, "七、修复汇总与安全对比", level=1)
    add_separator(doc)

    make_table(doc,
               ["安全维度", "修复前", "修复后"],
               [
                   ["外部实体解析", "允许 <!ENTITY SYSTEM 并自动读取文件", "拦截含 ENTITY/SYSTEM 的 XML"],
                   ["文件读取逻辑", "extract + open() 主动读取本地文件", "完全删除文件读取逻辑"],
                   ["路径遍历防御", "无防御，../ 可任意穿越", "拦截含 <!ENTITY 的 XML，路径穿越无效"],
                   ["错误信息", "暴露文件路径和具体错误", "统一提示不暴露细节"],
                   ["正常 XML 解析", "正常 XML 可解析", "正常 XML 可解析（保持不变）"],
                   ["JSON 输出", "正常输出", "正常输出"],
               ],
               col_widths=[3.5, 5, 5])

    # 八、修复后验证
    doc.add_page_break()
    add_heading_styled(doc, "八、修复后安全验证", level=1)
    add_separator(doc)

    make_table(doc,
               ["测试用例", "预期结果", "实际结果", "结论"],
               [
                   ["正常 XML 解析", "返回 JSON 结果", "正常解析并返回", "通过"],
                   ["XXE 读取 app.py", "拦截，不读取文件", "XML 包含不安全实体", "通过"],
                   ["XXE 读取 /etc/passwd", "拦截，不读取文件", "XML 包含不安全实体", "通过"],
                   ["XXE 路径遍历 ../../", "拦截，不读取文件", "XML 包含不安全实体", "通过"],
                   ["无效 XML 格式", "统一错误提示", "XML 格式错误", "通过"],
                   ["空 XML 提交", "提示输入数据", "请输入 XML 数据", "通过"],
               ],
               col_widths=[3.5, 3, 3.5, 1.5])

    add_para(doc, "验证结论：", bold=True, size=11, space_before=8)
    add_para(doc,
             "全部 4 种 XXE 利用方式已成功修复。正常 XML 解析功能完全保留，"
             "所有攻击 Payload 全部被拦截。修复后系统满足以下安全要求：\n"
             "（1）含 <!ENTITY 的 XML 被拒绝解析，外部实体无法被利用；\n"
             "（2）SYSTEM/PUBLIC 关键字被拦截，无法读取本地文件；\n"
             "（3）自动读取本地文件的高危逻辑已被删除；\n"
             "（4）错误信息统一化，不泄露细节。",
             first_indent=0.74)

    # 九、总结
    doc.add_page_break()
    add_heading_styled(doc, "九、实验总结与心得", level=1)
    add_separator(doc)

    add_para(doc,
             "通过本次 XXE 漏洞复现与修复实验，深入理解了 XML 外部实体注入漏洞的产生原理、"
             "多种利用手法及其对应的修复方案。",
             first_indent=0.74)

    add_para(doc, "核心认知：", bold=True, size=11)
    add_para(doc,
             "（1）XXE 漏洞的本质是 XML 解析器在处理用户提交的 XML 时，允许外部实体的加载和解析。"
             "即使 Python 的 ET.fromstring() 默认禁用外部实体，但如果后端代码主动提取 SYSTEM 路径"
             "并 open() 读取文件，就相当于重新实现了 XXE 攻击。",
             first_indent=0.74)

    add_para(doc,
             "（2）「不信任用户输入」是安全防御的第一原则。XML 文本完全由用户控制，其中可能包含"
             "恶意构造的 DTD 声明和外部实体引用。对 XML 输入进行关键字过滤是最直接的防御手段。",
             first_indent=0.74)

    add_para(doc,
             "（3）Python 的 xml.etree.ElementTree 在 Python 3.x 中默认不解析外部实体，"
             "但如果使用 lxml 或其他 XML 库，情况可能不同。在任何 XML 处理场景下，"
             "都应该显式禁用外部实体解析。",
             first_indent=0.74)

    add_para(doc,
             "通过本次实验深刻认识到：任何涉及用户提交 XML 的功能都应被视为高风险功能。"
             "防御 XXE 的最佳策略是：禁止外部实体 + 输入关键字过滤 + 最小化解析能力，三者缺一不可。",
             first_indent=0.74)

    # 结尾
    doc.add_paragraph()
    add_separator(doc)

    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run("报告完")
    set_run_font(run, size=11, color=(150, 150, 150))

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "本报告仅供安全教学与技术交流使用。报告中涉及的漏洞代码已全部修复。"
    )
    set_run_font(run, size=9, color=(180, 180, 180))

    doc.save(output_path)
    print(f"报告已生成：{output_path}")


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "XML导入功能_XXE外部实体注入漏洞复现及修复报告.docx")
    generate_report(out)
