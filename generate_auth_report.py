#!/usr/bin/env python3
"""
个人中心 & 充值模块 — 越权 / 金额 / 遍历漏洞复现及修复报告
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="微软雅黑", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return heading


def add_para(doc, text, bold=False, size=11, color=None, align=None,
             space_after=6, space_before=0, first_indent=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = Pt(20)
    if align:
        para.alignment = align
    if first_indent:
        para.paragraph_format.first_line_indent = Cm(first_indent)
    return para


def add_code_block(doc, code, indent=1.0):
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)


def add_separator(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def make_table(doc, headers, rows, col_widths=None, header_color="2F5496"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def generate_report(output_path):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(11)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ═══════════ 封面 ═══════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Flask 用户信息管理平台")
    set_run_font(run, size=26, bold=True, color=(47, 84, 150))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("个人中心 & 充值模块\n越权 / 金额 / 遍历漏洞复现及修复报告")
    set_run_font(run, size=18, bold=False, color=(89, 89, 89))

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━" * 40)
    set_run_font(run, size=12, color=(200, 200, 200))
    doc.add_paragraph()

    for text in [
        "实验性质：越权 / 金额逻辑 / ID 遍历漏洞复现与修复实训",
        "目标系统：Flask 用户信息管理系统",
        "文档版本：V1.0 — 终审版",
        "生成日期：2026 年 7 月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=12, color=(100, 100, 100))

    doc.add_page_break()

    # ═══════════ 目录 ═══════════
    add_heading_styled(doc, "目  录", level=1)
    doc.add_paragraph()
    toc_items = [
        "一、实验概述",
        "二、漏洞总览",
        "三、漏洞 1：水平 + 垂直越权查看任意用户 / 管理员隐私",
        "    3.1  复现内容",
        "    3.2  修复方案",
        "四、漏洞 2：越权篡改他人账户余额",
        "    4.1  复现内容",
        "    4.2  修复方案",
        "五、漏洞 3：负数金额非法扣款漏洞",
        "    5.1  复现内容",
        "    5.2  修复方案",
        "六、漏洞 4：单次充值无金额上限漏洞",
        "    6.1  复现内容",
        "    6.2  修复方案",
        "七、漏洞 5：user_id 可批量遍历致全站隐私泄露",
        "    7.1  复现内容",
        "    7.2  修复方案",
        "八、修复汇总与安全对比",
        "九、修复后安全验证",
        "十、实验总结与心得",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(item)
        set_run_font(run, size=11, bold=not item.startswith("    "))

    doc.add_page_break()

    # ═══════════ 一、实验概述 ═══════════
    add_heading_styled(doc, "一、实验概述", level=1)
    add_separator(doc)

    add_para(doc,
             "本次实验针对 Flask 用户信息管理平台中的个人中心与充值模块进行安全审计，"
             "共发现 5 项高危漏洞，涵盖水平/垂直越权、跨账号资金篡改、负数金额扣款、"
             "无上限充值、ID 批量遍历拖库等典型 Web 安全风险。",
             first_indent=0.74)
    add_para(doc,
             "实验流程包括：漏洞复现（Burp Suite 手工操作）→ 根因分析 → 安全加固 → "
             "修复后验证，完整覆盖从攻击到防御的全生命周期。",
             first_indent=0.74)

    # ═══════════ 二、漏洞总览 ═══════════
    add_heading_styled(doc, "二、漏洞总览", level=1)
    add_separator(doc)

    make_table(doc,
               ["编号", "漏洞名称", "漏洞类型", "危害等级"],
               [
                   ["VUL-AUTH-01", "水平 + 垂直越权查看任意用户隐私", "访问控制缺陷", "高危"],
                   ["VUL-AUTH-02", "越权篡改他人账户余额", "访问控制缺陷", "高危"],
                   ["VUL-AUTH-03", "负数金额非法扣款", "输入校验缺失", "高危"],
                   ["VUL-AUTH-04", "单次充值无金额上限", "业务逻辑缺陷", "中危"],
                   ["VUL-AUTH-05", "user_id 批量遍历致全站隐私泄露", "信息泄露", "高危"],
               ],
               col_widths=[2.5, 5.5, 3, 2])

    add_para(doc, "漏洞根因总结：", bold=True, size=11, space_before=6)
    add_para(doc, "（1）个人中心通过 URL 参数 user_id 定位用户，未验证归属关系；", first_indent=0.74)
    add_para(doc, "（2）充值接口信赖前端表单提交的 user_id，未强制从 session 读取；", first_indent=0.74)
    add_para(doc, "（3）amount 金额字段无正负校验、无上限校验；", first_indent=0.74)
    add_para(doc, "（4）无频率限制，ID 可被批量遍历爬取。", first_indent=0.74)

    # ═══════════════════════════════════════════
    # 漏洞 1
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "三、漏洞 1：水平 + 垂直越权查看任意用户 / 管理员隐私", level=1)
    add_separator(doc)

    add_heading_styled(doc, "3.1  复现内容", level=2)
    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "启动靶场，创建两个账号：普通用户 userA（user_id=2）、管理员 admin（user_id=1），记录两者邮箱、手机号、初始余额。",
        "打开 Burp Suite，浏览器配置代理，使用普通用户 userA 完成登录，保留有效登录 Session。",
        "点击导航栏「个人中心」，浏览器地址栏为 /profile?user_id=2，正常展示 userA 自身信息。",
        "不退出登录，直接修改 URL 参数：将 user_id=2 改为其他普通用户 ID 访问页面；再将 user_id 改为 1（管理员 ID），刷新页面。",
        "记录页面展示的全部用户隐私信息。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "修改为其他普通用户 ID 时，页面完整展示对方 ID、用户名、邮箱、手机号、余额，实现水平越权；"
             "修改为管理员 user_id=1 后，页面完整展示 admin 全部敏感隐私，实现垂直越权。"
             "后端未校验当前登录用户与查询 user_id 归属关系，低权限用户可随意查看全站用户、管理员数据，越权漏洞存在。",
             first_indent=0.74)

    add_para(doc, "【漏洞代码（修复前）】", bold=True, size=11, color=(192, 57, 43))
    add_code_block(doc,
        "# 修复前：直接从 URL 参数获取 user_id，未校验归属\n"
        'user_id = int(request.args.get("user_id", 0))\n'
        "username, user_data = get_user_by_id(user_id)\n"
        "# → 任意用户 ID 均可查询，无归属校验"
    )

    add_heading_styled(doc, "3.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "废弃 URL 参数 user_id，用户身份强制从服务端 session 读取。个人中心仅展示当前登录用户自身信息，"
             "拒绝接受外部可控的 user_id 参数，从根源上消除越权查询路径。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "session 中存储的 username 在登录时由服务端写入，客户端无法篡改。"
             "将数据查询条件与服务端会话绑定后，攻击者无法通过修改 URL 参数绕过身份校验。"
             "同时删除 _USER_ID_MAP、get_user_by_id() 等辅助函数，彻底关闭 user_id 查询通道。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        "# 修复后：用户身份强制从 session 读取\n"
        '@app.route("/profile")\n'
        "@login_required\n"
        "def profile():\n"
        '    username = session.get("username")    # 从 session 读取\n'
        "    user_data = USERS.get(username)\n"
        "    # URL 参数 user_id 被完全忽略"
    )

    # ═══════════════════════════════════════════
    # 漏洞 2
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "四、漏洞 2：越权篡改他人账户余额", level=1)
    add_separator(doc)

    add_heading_styled(doc, "4.1  复现内容", level=2)
    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "保持普通用户 userA 登录，进入个人中心页面，页面底部存在充值表单。",
        "表单输入正数 10，开启 Burp 拦截，点击充值按钮，捕获 POST /recharge 数据包。",
        "原始请求体携带隐藏字段 user_id=2&amount=10，将 user_id 篡改为管理员 user_id=1，amount 保持 10 不变。",
        "点击 Send 发送篡改后的请求，页面自动重定向至 /profile?user_id=1。",
        "查看管理员个人中心余额数值变化。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "发包后管理员账户余额直接增加 10；普通用户可通过抓包篡改 user_id，"
             "任意修改其他用户、管理员账户余额，跨账号资金越权漏洞存在。",
             first_indent=0.74)

    add_para(doc, "【漏洞代码（修复前）】", bold=True, size=11, color=(192, 57, 43))
    add_code_block(doc,
        "# 修复前：信赖前端表单提交的 user_id\n"
        'user_id = int(request.form.get("user_id", 0))\n'
        "amount = float(request.form.get(\"amount\", 0))\n"
        "username, user_data = get_user_by_id(user_id)  # user_id 可控"
    )

    add_heading_styled(doc, "4.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "充值目标用户强制从 session 读取，完全废弃表单提交的 user_id 参数，"
             "同时从充值表单中删除隐藏的 user_id 输入框。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "后端完全忽略前端提交的 user_id 字段，充值目标用户仅取自 session 中存储的登录用户名。"
             "攻击者即使篡改请求体中的任意参数，也无法改变充值的目标对象，"
             "因为后端代码中根本不存在从请求参数读取 user_id 的逻辑。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        "# 修复后：充值目标强制从 session 读取\n"
        '@app.route("/recharge", methods=["POST"])\n'
        "@login_required\n"
        "def recharge():\n"
        '    username = session.get("username")    # 强制从 session 读取\n'
        "    user_data = USERS.get(username)\n"
        "    # 表单 user_id 被完全忽略\n\n"
        "    # 表单中已删除隐藏 user_id 输入框"
    )

    # ═══════════════════════════════════════════
    # 漏洞 3
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "五、漏洞 3：负数金额非法扣款漏洞", level=1)
    add_separator(doc)

    add_heading_styled(doc, "5.1  复现内容", level=2)
    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "普通用户 userA 保持登录，进入充值页面，Burp 拦截 /recharge 提交请求。",
        "修改请求参数 amount=-999999999999，user_id 保留自身 2。",
        "Send 发送请求，页面重定向回 userA 个人中心，查看余额展示。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "充值请求提交成功，账户余额执行 balance + 负数逻辑，余额变为巨额负数；"
             "后端未校验金额正负，攻击者可恶意扣减自身或他人账户资金，资金逻辑漏洞存在。",
             first_indent=0.74)

    add_para(doc, "【漏洞代码（修复前）】", bold=True, size=11, color=(192, 57, 43))
    add_code_block(doc,
        "# 修复前：amount 无正负校验\n"
        'amount = float(request.form.get("amount", 0))\n'
        'user_data["balance"] = user_data.get("balance", 0) + amount  # 负数直接相加'
    )

    add_heading_styled(doc, "5.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "对 amount 实施三层校验：格式校验（纯数字正则）、正负校验（必须大于 0）、类型转换（float）。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "第一层：正则 ^[0-9]+(\\.[0-9]{1,2})?$ 拦截字母、负号、SQL 注入字符等非法输入。"
             "第二层：amount <= 0 拦截零和负数。"
             "三层依次过滤，任何一层不通过即返回明确错误提示，不执行充值操作。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        "# 修复后：amount 三层校验\n"
        'amount_str = request.form.get("amount", "").strip()\n\n'
        "# 第 1 层：非空校验\n"
        "if not amount_str:\n"
        '    return error("请输入充值金额")\n\n'
        "# 第 2 层：纯数字格式校验（拦截负号、字母、SQL注入）\n"
        'if not re.match(r"^[0-9]+(\\.[0-9]{1,2})?$", amount_str):\n'
        '    return error("充值金额不合法")\n\n'
        "# 第 3 层：正负校验\n"
        "amount = float(amount_str)\n"
        "if amount <= 0:\n"
        '    return error("充值金额必须大于 0")'
    )

    # ═══════════════════════════════════════════
    # 漏洞 4
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "六、漏洞 4：单次充值无金额上限漏洞", level=1)
    add_separator(doc)

    add_heading_styled(doc, "6.1  复现内容", level=2)
    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "userA 登录，拦截充值 POST 请求，修改 amount 为超大正数 999999999999。",
        "发送数据包，跳转个人中心查看余额。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "后端无单次充值上限限制，超大数值直接叠加至账户余额，用户余额数值异常膨胀，"
             "破坏平台财务数据规范，无上限充值漏洞存在。",
             first_indent=0.74)

    add_para(doc, "【漏洞代码（修复前）】", bold=True, size=11, color=(192, 57, 43))
    add_code_block(doc,
        "# 修复前：无充值上限\n"
        'user_data["balance"] = user_data.get("balance", 0) + amount  # 任意大额直接加'
    )

    add_heading_styled(doc, "6.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "设置硬性单次充值上限 MAX_RECHARGE_AMOUNT = 10000，amount 超过上限直接拦截。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "在正则校验和正负校验之后，增加上限校验。金额超过 10000 元即返回明确错误提示，"
             "不执行充值操作。上限值可在配置中统一调整，方便业务扩展。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        "# 配置：单次充值上限\n"
        "MAX_RECHARGE_AMOUNT = 10000\n\n"
        "# amount 通过格式校验和正负校验后\n"
        "if amount > MAX_RECHARGE_AMOUNT:\n"
        '    return error(f"单次充值金额不能超过 {MAX_RECHARGE_AMOUNT} 元")'
    )

    # ═══════════════════════════════════════════
    # 漏洞 5
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "七、漏洞 5：user_id 批量遍历致全站隐私泄露", level=1)
    add_separator(doc)

    add_heading_styled(doc, "7.1  复现内容", level=2)
    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "普通用户登录，Burp 捕获 GET /profile?user_id=1 数据包。",
        "使用 Burp Intruder 模块，对 user_id 参数设置数字爆破，payload 设置 1、2、3、4... 批量递增。",
        "批量发送请求，抓取每个响应页面返回的邮箱、手机号、余额数据。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "所有 user_id 均可正常查询，批量发包后可爬取全站所有用户的完整隐私信息，"
             "实现用户信息拖库，批量信息泄露高危漏洞存在。",
             first_indent=0.74)

    add_heading_styled(doc, "7.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "双重防御：一是废弃 user_id 查询逻辑（根本性修复），二是新增接口频率限制。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "第一重（根本）：个人中心不再接受 user_id 参数，数据仅从 session 获取，攻击者即使遍历 ID 也无法获取他人数据。"
             "第二重（增强）：新增内存级频率限制，单 IP 在 10 秒窗口内最多请求 10 次，"
             "超出限制后直接拒绝响应，从速率层面阻断批量遍历。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        "# 频率限制配置\n"
        "from collections import defaultdict\n"
        "_rate_limit_store = defaultdict(list)\n"
        "RATE_LIMIT_WINDOW = 10    # 10 秒窗口\n"
        "RATE_LIMIT_MAX = 10        # 窗口内最多 10 次\n\n"
        "def check_rate_limit(key_prefix):\n"
        '    key = f"{key_prefix}:{request.remote_addr}"\n'
        "    records = _rate_limit_store[key]\n"
        "    # 清理过期记录\n"
        '    _rate_limit_store[key] = [t for t in records if now - t < RATE_LIMIT_WINDOW]\n'
        "    if len(_rate_limit_store[key]) >= RATE_LIMIT_MAX:\n"
        "        return False\n"
        "    _rate_limit_store[key].append(now)\n"
        "    return True\n\n"
        "# 在 profile / recharge 入口处调用\n"
        "if not check_rate_limit(\"profile\"):\n"
        '    return error("请求过于频繁，请稍后重试")'
    )

    # ═══════════════════════════════════════════
    # 八、修复汇总
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "八、修复汇总与安全对比", level=1)
    add_separator(doc)

    make_table(doc,
               ["漏洞", "原代码缺陷", "修复方案", "修复技术"],
               [
                   ["越权查看他人资料",
                    "URL 参数 user_id 直接查库",
                    "从 session 读取，废弃 URL 参数",
                    "服务端会话绑定"],
                   ["越权修改他人余额",
                    "表单 user_id 直接定位目标",
                    "充值目标强制从 session 读取",
                    "服务端会话绑定"],
                   ["负数金额扣款",
                    "amount 直接 float 转换相加",
                    "正则 + 正负 + 格式三层校验",
                    "输入白名单校验"],
                   ["无上限充值",
                    "amount 无上限限制",
                    "MAX_RECHARGE_AMOUNT=10000",
                    "硬性上限拦截"],
                   ["ID 遍历拖库",
                    "user_id 可遍历查询",
                    "废弃 ID 查询 + 频率限制",
                    "双重防御"],
               ],
               col_widths=[3, 4, 4.5, 3])

    make_table(doc,
               ["安全维度", "修复前", "修复后"],
               [
                   ["用户身份来源", "URL / 表单参数 user_id（客户端可控）", "session（服务端写入，不可篡改）"],
                   ["跨用户访问", "低权限用户可查任意用户数据", "仅展示当前登录用户自身信息"],
                   ["金额正负校验", "无校验，负数直接扣款", "三层校验：格式>正负>上限"],
                   ["单次充值上限", "无上限，任意大额", "硬性上限 10000 元"],
                   ["批量遍历防御", "无限制，可批量爬取", "10 秒 / 10 次频率限制"],
                   ["user_id 注入", "字符串可能含 SQL 注入", "已废弃 user_id 查询逻辑"],
                   ["错误信息", '直接暴露"用户不存在"', "统一提示，不暴露信息"],
               ],
               col_widths=[3, 5.5, 6])

    # ═══════════════════════════════════════════
    # 九、修复后验证
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "九、修复后安全验证", level=1)
    add_separator(doc)

    make_table(doc,
               ["测试用例", "预期结果", "实际结果", "结论"],
               [
                   ["正常访问 /profile", "显示当前登录用户信息", "显示正确用户信息", "通过"],
                   ["访问 /profile?user_id=2 越权", "仍显示当前用户，忽略参数", "仍然显示当前用户", "通过"],
                   ["充值正数金额", "余额增加，充值成功", "余额正确增加", "通过"],
                   ["充值负数金额", "拦截，返回错误提示", '显示"充值金额不合法"', "通过"],
                   ["充值零元", "拦截，返回错误提示", '显示"充值金额必须大于 0"', "通过"],
                   ["充值超上限 99999", "拦截，返回错误提示", '显示"不能超过 10000"', "通过"],
                   ["充值 SQL 注入字符", "拦截，返回错误提示", '显示"充值金额不合法"', "通过"],
                   ["高频请求越权遍历", "拦截频率过高请求", "返回频率限制提示", "通过"],
                   ["充值表单含 user_id", "不存在 user_id 字段", "表单无 user_id", "通过"],
               ],
               col_widths=[4, 3.5, 3.5, 1.5])

    add_para(doc, "验证结论：", bold=True, size=11, space_before=8)
    add_para(doc,
             "全部 5 项漏洞已成功修复。正常业务功能完全保留，攻击向量全部被拦截。"
             "用户身份强制从 session 读取，URL 和表单参数 user_id 被完全忽略；"
             "金额实施三层校验，负数、超上限、非法格式全部被拒；"
             "频率限制有效阻断批量遍历攻击。",
             first_indent=0.74)

    # ═══════════════════════════════════════════
    # 十、总结
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "十、实验总结与心得", level=1)
    add_separator(doc)

    add_para(doc,
             "通过本次个人中心与充值模块漏洞复现与修复实验，系统掌握了 Web 应用中常见的"
             "越权漏洞、资金逻辑漏洞、ID 遍历漏洞的产生原理与标准化修复方案。",
             first_indent=0.74)

    add_para(doc, "核心认知：", bold=True, size=11)
    add_para(doc,
             "（1）用户身份识别必须依赖服务端 session，任何由客户端传入的用户标识（URL 参数、"
             "表单字段、Cookie 中的 ID）都是不可信的，必须与服务端会话绑定校验。",
             first_indent=0.74)
    add_para(doc,
             "（2）涉及资金操作的接口必须实施多重校验：格式校验确保数据类型正确、"
             "逻辑校验确保数值范围合理、权限校验确保操作者身份合法。",
             first_indent=0.74)
    add_para(doc,
             "（3）批量遍历漏洞的修复应遵循「根本修复 + 增强防御」的双重策略，"
             "既要从架构上消除遍历路径，也要从速率上阻断批量请求。"
             "单一措施（如仅加频率限制）可能被绕过，双重防御更为可靠。",
             first_indent=0.74)

    add_para(doc,
             "通过本次实验深刻认识到，Web 安全中最危险的漏洞往往不是技术门槛最高的漏洞，"
             "而是开发时最容易忽视的「默认信任客户端输入」的习惯。"
             "后端代码必须对每一个外部输入都持不信任态度，严格验证、层层过滤。",
             first_indent=0.74)

    # 结尾
    doc.add_paragraph()
    add_separator(doc)

    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run("— 报告完 —")
    set_run_font(run, size=11, color=(150, 150, 150))

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "本报告仅供安全教学与技术交流使用。报告中涉及的漏洞代码已全部修复，\n"
        "请勿将存在漏洞的代码版本部署到任何生产环境。"
    )
    set_run_font(run, size=9, color=(180, 180, 180))

    doc.save(output_path)
    print(f"报告已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "个人中心与充值模块_越权金额遍历漏洞复现及修复报告.docx")
    generate_report(out)
