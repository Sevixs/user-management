#!/usr/bin/env python3
"""
修改密码模块 — 越权 / CSRF / 弱密码 / 无原码校验漏洞复现及修复报告
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="微软雅黑", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return heading


def add_para(doc, text, bold=False, size=11, color=None, align=None,
             space_after=6, space_before=0, first_indent=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = Pt(20)
    if align:
        para.alignment = align
    if first_indent:
        para.paragraph_format.first_line_indent = Cm(first_indent)
    return para


def add_code_block(doc, code, indent=1.0):
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)


def add_separator(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def make_table(doc, headers, rows, col_widths=None, header_color="1A5276"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def generate_report(output_path):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(11)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ═══════════ 封面 ═══════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Flask 用户信息管理平台")
    set_run_font(run, size=26, bold=True, color=(26, 82, 118))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("修改密码模块\n越权 / CSRF / 弱密码 / 无原码校验\n漏洞复现及修复报告")
    set_run_font(run, size=18, bold=False, color=(89, 89, 89))

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("-" * 40)
    set_run_font(run, size=12, color=(200, 200, 200))
    doc.add_paragraph()

    for text in [
        "实验性质：越权 / CSRF / 弱密码漏洞复现与修复实训",
        "目标系统：Flask 用户信息管理系统（/change-password 路由）",
        "文档版本：V1.0 终审版",
        "生成日期：2026 年 7 月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=12, color=(100, 100, 100))

    doc.add_page_break()

    # ═══════════ 目录 ═══════════
    add_heading_styled(doc, "目  录", level=1)
    doc.add_paragraph()
    toc_items = [
        "一、实验概述",
        "二、漏洞总览",
        "三、漏洞 1：越权篡改管理员密码",
        "    3.1  复现内容",
        "    3.2  修复方案",
        "四、漏洞 2：无需原始密码即可修改密码",
        "    4.1  复现内容",
        "    4.2  修复方案",
        "五、漏洞 3：CSRF 跨站伪造修改密码",
        "    5.1  复现内容",
        "    5.2  修复方案",
        "六、漏洞 4：允许设置空密码 / 极简弱密码",
        "    6.1  复现内容",
        "    6.2  修复方案",
        "七、修复方案汇总与安全对比",
        "八、修复后安全验证",
        "九、实验总结与心得",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(item)
        set_run_font(run, size=11, bold=not item.startswith("    "))

    doc.add_page_break()

    # ═══════════ 一、实验概述 ═══════════
    add_heading_styled(doc, "一、实验概述", level=1)
    add_separator(doc)

    add_para(doc,
             "本次实验针对 Flask 用户信息管理平台中的密码修改模块 /change-password 进行安全审计。"
             "该路由用于已登录用户修改自身密码，但由于缺乏必要的安全校验，导致存在越权篡改、"
             "无原码校验、CSRF 跨站伪造、弱密码等 4 项高危漏洞。",
             first_indent=0.74)
    add_para(doc,
             "实验流程包括：漏洞复现（Burp Suite 手工操作）-> 根因分析 -> 安全加固 -> "
             "修复后验证，完整覆盖从攻击到防御的全生命周期。",
             first_indent=0.74)

    # ═══════════ 二、漏洞总览 ═══════════
    add_heading_styled(doc, "二、漏洞总览", level=1)
    add_separator(doc)

    make_table(doc,
               ["编号", "漏洞名称", "漏洞类型", "危害等级"],
               [
                   ["VUL-PWD-01", "越权篡改管理员密码", "访问控制缺陷", "高危"],
                   ["VUL-PWD-02", "无需原始密码即可修改密码", "身份校验缺失", "高危"],
                   ["VUL-PWD-03", "CSRF 跨站伪造修改密码", "跨站请求伪造", "高危"],
                   ["VUL-PWD-04", "允许设置空密码 / 极简弱密码", "密码策略缺失", "中危"],
               ],
               col_widths=[2.5, 5.5, 3.5, 1.5])

    add_para(doc, "漏洞根因总结：", bold=True, size=11, space_before=6)
    add_para(doc, "（1）修改密码的操作目标用户由前端表单 username 控制，未与服务端 session 绑定；", first_indent=0.74)
    add_para(doc, "（2）修改密码接口未校验用户原始密码，仅需新密码即可覆盖；", first_indent=0.74)
    add_para(doc, "（3）修改密码接口无 CSRF Token 校验，可被跨站伪造；", first_indent=0.74)
    add_para(doc, "（4）新密码无格式复杂度校验，空密码和弱密码均可通过。", first_indent=0.74)

    add_para(doc, "漏洞代码片段（修复前）：", bold=True, size=10, color=(192, 57, 43))
    add_code_block(doc,
        '# 修复前：无过滤、无校验、无权限\n'
        '@app.route("/change-password", methods=["POST"])\n'
        "@login_required\n"
        "def change_password():\n"
        '    username = request.form.get("username", "")    # 前端可控\n'
        '    new_password = request.form.get("new_password", "")\n'
        "    if username in USERS:\n"
        '        USERS[username]["password"] = new_password  # 直接覆盖，无校验'
    )

    # ═══════════════════════════════════════════
    # 漏洞 1
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "三、漏洞 1：越权篡改管理员密码", level=1)
    add_separator(doc)

    add_heading_styled(doc, "3.1  复现内容", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "注册并登录普通用户 userA，进入个人中心 /profile 页面。",
        "页面底部找到修改密码表单，开启 Burp Suite 拦截表单提交请求。",
        "表单输入任意新密码，点击修改密码，捕获 POST /change-password 数据包。",
        "请求体内存在隐藏字段 username=userA，将参数篡改为 username=admin，new_password 保留自定义新密码。",
        "点击 Send 发送篡改后的请求，页面自动重定向至 /profile?user_id=1。",
        "退出当前 userA 登录，打开登录页面，使用新密码尝试登录管理员 admin 账号。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "数据包提交成功，管理员数据库内密码被直接替换为攻击者设置的新密码；"
             "使用新密码可正常登录管理员账号，水平 + 垂直越权篡改密码漏洞存在。",
             first_indent=0.74)

    add_para(doc, "攻击原理示意图：", bold=True, size=11)
    add_code_block(doc,
        "原始请求: username=userA&new_password=hack123\n"
        "篡改后:   username=admin&new_password=hack123\n"
        "后端处理: USERS['admin']['password'] = 'hack123'\n"
        "结果: 攻击者以 userA 身份篡改了 admin 的密码"
    )

    add_heading_styled(doc, "3.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "完全废弃表单提交的 username 参数，修改密码的操作对象强制从 session 读取。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "后端代码中根本不存在从 request.form 读取 username 的逻辑。"
             "目标用户仅通过 session.get('username') 获取，session 由服务端在登录时写入，"
             "客户端无法篡改。即使攻击者在请求体中添加任意 username 参数，后端也不会使用。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        '# 修复后：目标用户强制从 session 读取\n'
        'session_username = session.get("username")\n'
        "user_data = USERS.get(session_username)\n"
        "# 完全忽略 request.form.get('username')\n\n"
        "# profile.html 中已删除 username 隐藏字段\n"
        '# 表单不再携带 username 参数'
    )

    # ═══════════════════════════════════════════
    # 漏洞 2
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "四、漏洞 2：无需原始密码即可修改密码", level=1)
    add_separator(doc)

    add_heading_styled(doc, "4.1  复现内容", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "登录 userA 账号，进入修改密码表单。",
        "仅填写新密码，不存在原密码输入框，直接提交表单。",
        "查看个人中心页面，确认密码修改成功。",
        "退出登录，使用新密码登录 userA。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "无需输入、校验原有账号密码，直接覆盖密码。"
             "攻击者仅需知道用户名就能重置密码，无原密码校验漏洞存在。",
             first_indent=0.74)

    add_heading_styled(doc, "4.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "新增 old_password 输入框，后端使用 Werkzeug check_password_hash 校验原始密码哈希。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "用户密码以 scrypt 哈希存储在 USERS 字典中。修改密码时，用户必须提供原密码，"
             "后端通过 check_password_hash(user_data['password'], old_password) 比对哈希值。"
             "即使数据库泄露，攻击者也需先破解原密码哈希才能修改密码。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        '# 修复后：新增原密码哈希校验\n'
        'old_password = request.form.get("old_password", "")\n'
        "if not check_password_hash(user_data[\"password\"], old_password):\n"
        '    return error("原密码错误")\n\n'
        '# 通过后才允许设置新密码\n'
        'user_data["password"] = generate_password_hash(new_password)'
    )

    # ═══════════════════════════════════════════
    # 漏洞 3
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "五、漏洞 3：CSRF 跨站伪造修改密码", level=1)
    add_separator(doc)

    add_heading_styled(doc, "5.1  复现内容", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "搭建简易恶意 HTML 页面，包含自动提交到 /change-password 的表单。",
        "保持浏览器内 userA 处于登录状态（保留有效 session）。",
        "本地打开恶意 HTML 页面，点击页面中的按钮触发表单自动提交。",
        "返回网站登录页面，使用攻击者预设的密码登录 userA。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "恶意 HTML 页面代码：", bold=True, size=10, color=(192, 57, 43))
    add_code_block(doc,
        '<form action="http://127.0.0.1:5000/change-password" method="POST">\n'
        '    <input type="hidden" name="username" value="userA">\n'
        '    <input type="hidden" name="new_password" value="hack123456">\n'
        '    <button type="submit">领取福利</button>\n'
        "</form>"
    )

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "跨域页面成功发送改密请求，用户不知情情况下密码被篡改，"
             "CSRF 跨站请求伪造漏洞存在。",
             first_indent=0.74)

    add_heading_styled(doc, "5.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "将 change-password 纳入 CSRF Token 校验范围（从排除名单中移除），"
             "所有 POST 请求必须携带有效的 csrf_token。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "CSRF Token 由 secrets.token_hex(32) 生成，绑定在用户 session 中。"
             "攻击者构造的跨站页面无法获取目标用户的 session Token，因此无法通过校验。"
             "同时表单页面增加隐藏的 csrf_token 字段，正常用户提交时 Token 由模板自动注入。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        '# CSRF 排除名单中移除 change_password\n'
        'if request.method == "POST" \\\n'
        '   and request.endpoint not in ("static", "serve_upload"):\n'
        "    # change_password 不再排除，需校验 Token\n\n"
        "    token = request.form.get(\"csrf_token\")\n"
        "    stored = session.get(\"csrf_token\")\n"
        "    if not token or not stored or token != stored:\n"
        '        abort(400, "CSRF token 缺失或无效")\n\n'
        '<!-- profile.html 表单中增加 -->\n'
        '<input type="hidden" name="csrf_token" value="{{ csrf_token }}">'
    )

    # ═══════════════════════════════════════════
    # 漏洞 4
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "六、漏洞 4：允许设置空密码 / 极简弱密码", level=1)
    add_separator(doc)

    add_heading_styled(doc, "6.1  复现内容", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "登录 userA，进入修改密码表单。",
        "新密码输入框填写空内容 / 填写 123456 简单 6 位数字密码。",
        "提交修改密码表单。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "后端无校验，空密码、6 位纯数字弱密码均可成功写入数据库，弱口令风险漏洞存在。",
             first_indent=0.74)

    add_heading_styled(doc, "6.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "实施 4 项密码复杂度校验：长度不少于 8 位 + 含大写字母 + 含小写字母 + 含数字。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "使用正则和 len() 函数对 new_password 进行逐项检查："
             "len >= 8 确保基础强度；re.search(r'[a-z]', pwd) 确保含小写；"
             "re.search(r'[A-Z]', pwd) 确保含大写；re.search(r'[0-9]', pwd) 确保含数字。"
             "四项校验独立进行，任何一项不满足即返回明确错误提示，拒绝修改。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        '# 修复后：4 项密码复杂度校验\n'
        "if len(new_password) < 8:\n"
        '    return error("密码长度不能少于 8 位")\n'
        "if not re.search(r\"[a-z]\", new_password):\n"
        '    return error("密码必须包含小写字母")\n'
        "if not re.search(r\"[A-Z]\", new_password):\n"
        '    return error("密码必须包含大写字母")\n'
        "if not re.search(r\"[0-9]\", new_password):\n"
        '    return error("密码必须包含数字")\n\n'
        "# 全部通过后才执行密码更新"
    )

    # ═══════════════════════════════════════════
    # 七、修复汇总
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "七、修复方案汇总与安全对比", level=1)
    add_separator(doc)

    make_table(doc,
               ["漏洞", "攻击手法", "修复前", "修复后"],
               [
                   ["越权篡改密码", "表单username改为admin", "直接覆盖管理员密码", "session读取，忽略表单"],
                   ["无原码校验", "直接输入新密码提交", "无需原密码即可重置", "check_password_hash校验"],
                   ["CSRF伪造", "跨站表单自动提交", "无Token校验，直接执行", "CSRF Token 全局校验"],
                   ["弱密码", "空/123456/纯数字", "直接写入，无校验", "8位+大小写+数字"],
               ],
               col_widths=[2.5, 3, 3.5, 4.5])

    make_table(doc,
               ["安全维度", "修复前", "修复后"],
               [
                   ["操作对象来源", "表单username（客户端可控）", "session（服务端写入）"],
                   ["跨账号修改", "可修改任意用户密码", "仅可修改自身密码"],
                   ["原密码校验", "无原码输入框和校验", "check_password_hash哈希比对"],
                   ["CSRF 防护", "排除在CSRF校验之外", "纳入CSRF Token全覆盖"],
                   ["密码长度", "无限制，可为空", ">= 8 位"],
                   ["密码复杂度", "无限制", "大写 + 小写 + 数字"],
                   ["密码存储", "明文直接存储", "generate_password_hash加密"],
                   ["错误提示", "无提示/直接成功", "具体错误原因提示"],
               ],
               col_widths=[3, 4.5, 5.5])

    # ═══════════════════════════════════════════
    # 八、修复后验证
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "八、修复后安全验证", level=1)
    add_separator(doc)

    make_table(doc,
               ["测试用例", "预期结果", "实际结果", "结论"],
               [
                   ["正常改密(原码正确+密码合规)", "修改成功", "密码修改成功", "通过"],
                   ["表单username改为admin", "仍修改自身密码", "仅修改当前用户", "通过"],
                   ["原密码错误提交", "拦截，错误提示", "显示原密码错误", "通过"],
                   ["CSRF Token 缺失", "400 拦截", "400 错误响应", "通过"],
                   ["CSRF Token 错误", "400 拦截", "400 错误响应", "通过"],
                   ["新密码长度<8", "拦截，错误提示", "提示密码长度不够", "通过"],
                   ["新密码无大写字母", "拦截，错误提示", "提示需含大写字母", "通过"],
                   ["新密码无小写字母", "拦截，错误提示", "提示需含小写字母", "通过"],
                   ["新密码无数字", "拦截，错误提示", "提示需含数字", "通过"],
                   ["两次密码不一致", "拦截，错误提示", "提示密码不一致", "通过"],
                   ["新密码为空字符串", "拦截，错误提示", "提示填写完整信息", "通过"],
                   ["profile表单无username字段", "不存在该字段", "确实无此字段", "通过"],
               ],
               col_widths=[4, 3, 3.5, 1.5])

    add_para(doc, "验证结论：", bold=True, size=11, space_before=8)
    add_para(doc,
             "全部 4 项漏洞已成功修复。正常改密功能完全保留，所有攻击向量全部被拦截。"
             "修复后系统满足以下安全要求：\n"
             "（1）修改密码操作对象强制从 session 读取，表单 username 被完全忽略；\n"
             "（2）原密码通过哈希校验，无法绕过；\n"
             "（3）CSRF Token 全覆盖，跨站伪造请求被拒绝；\n"
             "（4）密码复杂度 4 项校验，弱密码和空密码无法通过；\n"
             "（5）密码以 Werkzeug scrypt 哈希加密存储，不再明文保存。",
             first_indent=0.74)

    # ═══════════════════════════════════════════
    # 九、总结
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "九、实验总结与心得", level=1)
    add_separator(doc)

    add_para(doc,
             "通过本次密码修改模块漏洞复现与修复实验，深入理解了敏感操作中必须实施的"
             "多重安全校验机制。",
             first_indent=0.74)

    add_para(doc, "核心认知：", bold=True, size=11)
    add_para(doc,
             "（1）敏感操作的「操作对象」必须从服务端 session 获取。任何由前端传入的用户标识"
             "（username、user_id 等）都是不可信的，前端数据仅适用于非敏感信息的展示和查询。",
             first_indent=0.74)
    add_para(doc,
             "（2）「知道你是谁」和「证明你是你」是两回事。登录状态（session）仅证明用户已通过认证，"
             "但敏感操作（改密、转账等）还需额外的身份确认，如原密码校验或二次验证。",
             first_indent=0.74)
    add_para(doc,
             "（3）CSRF Token 是防御跨站请求伪造的标准方案。所有涉及数据变更的 POST 请求"
             "都应纳入 CSRF 校验范围，不应存在「豁免」端点。",
             first_indent=0.74)
    add_para(doc,
             "（4）密码策略是账户安全的第一道防线。空密码和弱密码使得其他所有安全措施形同虚设，"
             "因为攻击者无需绕过任何校验即可直接猜解登录。复杂度校验是密码策略的基础要求。",
             first_indent=0.74)

    add_para(doc,
             "通过本次实验深刻认识到，密码修改功能是 Web 应用中风险最高的功能之一，"
             "必须从身份校验、权限控制、跨站防护、密码策略四个维度进行全面加固，"
             "任何一维度的缺失都可能导致整个账户体系被攻破。",
             first_indent=0.74)

    # 结尾
    doc.add_paragraph()
    add_separator(doc)

    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run("报告完")
    set_run_font(run, size=11, color=(150, 150, 150))

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "本报告仅供安全教学与技术交流使用。报告中涉及的漏洞代码已全部修复，\n"
        "请勿将存在漏洞的代码版本部署到任何生产环境。"
    )
    set_run_font(run, size=9, color=(180, 180, 180))

    doc.save(output_path)
    print(f"报告已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "修改密码模块_越权CSRF弱密码漏洞复现及修复报告.docx")
    generate_report(out)
