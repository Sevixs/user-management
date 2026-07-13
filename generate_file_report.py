#!/usr/bin/env python3
"""
动态页面文件包含漏洞 — 路径遍历 / 目录穿越 / 任意文件读取
复现及修复报告
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="微软雅黑", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return heading


def add_para(doc, text, bold=False, size=11, color=None, align=None,
             space_after=6, space_before=0, first_indent=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = Pt(20)
    if align:
        para.alignment = align
    if first_indent:
        para.paragraph_format.first_line_indent = Cm(first_indent)
    return para


def add_code_block(doc, code, indent=1.0):
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, color=(80, 80, 80))
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)


def add_separator(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def make_table(doc, headers, rows, col_widths=None, header_color="8E44AD"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
            p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def generate_report(output_path):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(11)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ═══════════ 封面 ═══════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Flask 用户信息管理平台")
    set_run_font(run, size=26, bold=True, color=(142, 68, 173))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("动态页面文件包含漏洞\n路径遍历 / 目录穿越 / 任意文件读取\n复现及修复报告")
    set_run_font(run, size=18, bold=False, color=(89, 89, 89))

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━" * 40)
    set_run_font(run, size=12, color=(200, 200, 200))
    doc.add_paragraph()

    for text in [
        "实验性质：文件包含 / 路径遍历漏洞复现与修复实训",
        "目标系统：Flask 用户信息管理系统（/page 路由）",
        "文档版本：V1.0 — 终审版",
        "生成日期：2026 年 7 月",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=12, color=(100, 100, 100))

    doc.add_page_break()

    # ═══════════ 目录 ═══════════
    add_heading_styled(doc, "目  录", level=1)
    doc.add_paragraph()
    toc_items = [
        "一、实验概述",
        "二、漏洞总览",
        "三、漏洞 1：路径遍历读取项目源码 app.py",
        "    3.1  复现内容",
        "    3.2  修复方案",
        "四、漏洞 2：利用自动补 .html 读取模板 base.html",
        "    4.1  复现内容",
        "    4.2  修复方案",
        "五、漏洞 3：读取无后缀敏感配置文件 config",
        "    5.1  复现内容",
        "    5.2  修复方案",
        "六、漏洞 4：读取服务器系统敏感文件",
        "    6.1  复现内容（Linux / Windows）",
        "    6.2  修复方案",
        "七、修复方案汇总与安全对比",
        "八、修复后安全验证",
        "九、实验总结与心得",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(item)
        set_run_font(run, size=11, bold=not item.startswith("    "))

    doc.add_page_break()

    # ═══════════ 一、实验概述 ═══════════
    add_heading_styled(doc, "一、实验概述", level=1)
    add_separator(doc)

    add_para(doc,
             "本次实验针对 Flask 用户信息管理平台中的动态页面加载路由 /page 进行安全审计。"
             "该路由用于加载 pages/ 目录下的静态帮助页面，但由于直接拼接用户输入的 name 参数"
             "构造文件路径且未做任何过滤，导致存在严重的文件包含与路径遍历漏洞。",
             first_indent=0.74)

    add_para(doc,
             "本次实验共发现 4 项高危漏洞，可读取项目源码、前端模板、服务器配置文件、"
             "操作系统敏感文件等。攻击者无需登录即可利用，危害极大。",
             first_indent=0.74)

    add_para(doc,
             "实验流程包括：漏洞复现（浏览器手工操作）-> 根因分析 -> 安全加固 -> "
             "修复后验证，完整覆盖从攻击到防御的全生命周期。",
             first_indent=0.74)

    # ═══════════ 二、漏洞总览 ═══════════
    add_heading_styled(doc, "二、漏洞总览", level=1)
    add_separator(doc)

    make_table(doc,
               ["编号", "漏洞名称", "漏洞类型", "危害等级"],
               [
                   ["VUL-FILE-01", "路径遍历读取项目源码 app.py", "路径遍历", "高危"],
                   ["VUL-FILE-02", "利用自动补 .html 读取模板文件", "路径遍历 + 文件包含", "高危"],
                   ["VUL-FILE-03", "读取无后缀敏感配置文件", "任意文件读取", "高危"],
                   ["VUL-FILE-04", "读取服务器系统敏感文件", "路径遍历 + 文件包含", "高危"],
               ],
               col_widths=[2.5, 6, 3, 1.5])

    add_para(doc, "漏洞根因总结：", bold=True, size=11, space_before=6)
    add_para(doc, "（1）未对 URL 参数 name 做任何过滤，直接拼接路径；", first_indent=0.74)
    add_para(doc, "（2）未拦截 ../、/ 等目录穿越字符；", first_indent=0.74)
    add_para(doc, "（3）未做路径规范化与目录归属校验；", first_indent=0.74)
    add_para(doc, "（4）无后缀名限制，任何类型文件均可读取；", first_indent=0.74)
    add_para(doc, "（5）自动补 .html 逻辑扩大了攻击面，可读取更多文件。", first_indent=0.74)

    add_para(doc, "漏洞代码片段（修复前）：", bold=True, size=10, color=(192, 57, 43))
    add_code_block(doc,
        '# 修复前：无过滤、无校验、无限制\n'
        '@app.route("/page")\n'
        "def dynamic_page():\n"
        '    name = request.args.get("name", "")\n'
        '    page_path = os.path.join("pages", name)     # 直接拼接\n'
        "    if os.path.isfile(page_path):\n"
        "        with open(page_path, \"r\") as f:\n"
        "            page_content = f.read()\n"
        "    # 还会自动补 .html 后缀再次尝试"
    )

    # ═══════════════════════════════════════════
    # 漏洞 1
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "三、漏洞 1：路径遍历读取项目源码 app.py", level=1)
    add_separator(doc)

    add_heading_styled(doc, "3.1  复现内容", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "启动 Flask 靶场，打开浏览器访问网站首页。",
        "不点击页面自带帮助链接，直接修改浏览器地址栏 URL。",
        '输入恶意地址：http://127.0.0.1:5000/page?name=../app.py，回车访问。',
        "观察页面展示内容。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "后端拼接路径 pages/../app.py，文件真实存在，页面完整输出 app.py 全部后端代码，"
             "包括数据库配置、密钥、路由逻辑、业务代码等，源代码泄露，路径遍历漏洞可利用。",
             first_indent=0.74)

    add_para(doc, "攻击原理示意图：", bold=True, size=11)
    add_code_block(doc,
        "URL 参数: name=../app.py\n"
        "拼接路径: pages/../app.py\n"
        "规范化后: ./app.py  ← 跳出了 pages 目录\n"
        "页面输出: app.py 全部源代码"
    )

    add_heading_styled(doc, "3.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "四层防御：黑名单拦截路径穿越字符 + 白名单校验合法页面名 + "
             "绝对路径规范化 + 目录归属强校验。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "第一层（黑名单）：检查 name 中是否包含 ..、/、\\、~ 等路径穿越字符，"
             "有则直接返回「页面不存在」，阻止路径拼接前的穿越尝试。\n"
             "第二层（白名单）：正则 ^[a-zA-Z0-9_\\-\\.]+$ 仅允许合法页面名字符，"
             "拦截一切特殊符号。\n"
             "第三层（路径规范化）：使用 os.path.abspath() 将拼接后的路径转换为绝对路径，"
             "消除 ../ 的目录跳跃效果。\n"
             "第四层（目录归属）：强制校验绝对路径是否以 pages 目录绝对路径开头，"
             "不在 pages 目录内的文件一律拒绝读取。",
             first_indent=0.74)

    add_para(doc, "【修复后代码】", bold=True, size=11, color=(39, 174, 96))
    add_code_block(doc,
        '# 安全加固版：四层防御\n'
        '# 第 1 层：黑名单拦截\n'
        'if ".." in name or "/" in name or "\\\\" in name or "~" in name:\n'
        '    return error("页面不存在")\n\n'
        '# 第 2 层：白名单校验\n'
        'if not re.match(r"^[a-zA-Z0-9_\\\\-.]+$", name):\n'
        '    return error("页面不存在")\n\n'
        '# 第 3 层：绝对路径规范化\n'
        'real_path = os.path.abspath(os.path.join(PAGES_DIR, name))\n\n'
        '# 第 4 层：目录归属校验\n'
        'if not real_path.startswith(os.path.abspath(PAGES_DIR) + os.sep):\n'
        '    return error("页面不存在")\n'
        'if not real_path.endswith(".html"):\n'
        '    return error("页面不存在")'
    )

    # ═══════════════════════════════════════════
    # 漏洞 2
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "四、漏洞 2：利用自动补 .html 读取模板 base.html", level=1)
    add_separator(doc)

    add_heading_styled(doc, "4.1  复现内容", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "保持网站首页打开，修改地址栏 URL。",
        "输入 Payload：http://127.0.0.1:5000/page?name=../templates/base，回车访问。",
        "后端逻辑：先读取 pages/../templates/base（无此文件），自动追加 .html，再读取 pages/../templates/base.html。",
        "查看页面返回内容。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "页面完整展示 base.html 前端模板源码，页面内所有导航链接、隐藏接口全部泄露。"
             "攻击者可遍历 templates 目录下全部页面，获取所有前端代码、表单字段、接口路径等敏感信息。",
             first_indent=0.74)

    add_para(doc, "攻击原理示意图：", bold=True, size=11)
    add_code_block(doc,
        "URL 参数: name=../templates/base\n"
        "第 1 次尝试: pages/../templates/base  → 文件不存在\n"
        "第 2 次尝试: pages/../templates/base.html  → 文件存在！\n"
        "页面输出: base.html 完整模板源码（导航/接口/表单字段全部泄露）"
    )

    add_heading_styled(doc, "4.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "自动补 .html 逻辑执行后同样执行目录归属校验和后缀校验，"
             "文件不在 pages 目录内直接拦截。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "自动补 .html 后缀发生在路径穿越字符拦截之后，此时 name 中的 ../ 已被拦截，"
             "根本不会进入文件读取逻辑。即使黑名单被绕过，路径规范化后的绝对路径"
             "也会被目录归属校验拦截——因为 templates 目录不在 pages 目录下。",
             first_indent=0.74)

    add_para(doc, "防护链路：", bold=True, size=11)
    add_code_block(doc,
        "name=../templates/base\n"
        "-> 第 1 层黑名单: 检测到 '..' → 拦截，返回「页面不存在」\n"
        "-> 后续校验不再执行\n\n"
        "即使绕过黑名单:\n"
        "-> 第 3 层规范化: real_path = /project/templates/base.html\n"
        "-> 第 4 层目录归属: /project/templates/ 不以 /project/pages/ 开头 → 拦截"
    )

    # ═══════════════════════════════════════════
    # 漏洞 3
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "五、漏洞 3：读取无后缀敏感配置文件 config", level=1)
    add_separator(doc)

    add_heading_styled(doc, "5.1  复现内容", level=2)

    add_para(doc, "【复现步骤】", bold=True, size=11)
    steps = [
        "回到网站首页，修改 URL 参数。",
        "访问地址：http://127.0.0.1:5000/page?name=../config。",
        "页面加载完成后查看展示内容。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【复现结果】", bold=True, size=11)
    add_para(doc,
             "后端直接读取 pages/../config，文件存在且不会自动拼接 .html（因文件已存在），"
             "页面输出配置文件内密钥、数据库地址等敏感信息，配置泄露漏洞存在。",
             first_indent=0.74)

    add_para(doc, "攻击原理示意图：", bold=True, size=11)
    add_code_block(doc,
        "URL 参数: name=../config\n"
        "拼接路径: pages/../config\n"
        "文件存在: 直接读取，不拼接 .html\n"
        "页面输出: 配置文件全文（密钥/数据库地址/API Key）"
    )

    add_heading_styled(doc, "5.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "强制要求文件必须以 .html 结尾，无后缀文件和非 html 文件一律拒绝读取。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "无论文件是否存在、是否有后缀、是否在 pages 目录内，最终读取前执行 .endswith('.html') 校验。"
             "config 文件无 .html 后缀，直接返回「页面不存在」。"
             "即使存在 pages/config.html 这样的文件，也仅能读取 pages 目录内的 html 页面，无法越界。",
             first_indent=0.74)

    add_para(doc, "防护链路：", bold=True, size=11)
    add_code_block(doc,
        "name=../config\n"
        "-> 第 1 层黑名单: 检测到 '..' → 拦截\n\n"
        "即使绕过黑名单:\n"
        "-> 拼接路径 = pages/../config\n"
        "-> 规范化路径 = /project/config\n"
        "-> 第 4 层后缀校验: 不以 .html 结尾 → 拦截\n\n"
        "name=../config.json\n"
        "-> 后缀校验: 不以 .html 结尾 → 拦截"
    )

    # ═══════════════════════════════════════════
    # 漏洞 4
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "六、漏洞 4：读取服务器系统敏感文件", level=1)
    add_separator(doc)

    add_heading_styled(doc, "6.1  复现内容", level=2)

    add_para(doc, "【Linux 系统复现步骤】", bold=True, size=11)
    steps = [
        "首页地址栏输入：http://127.0.0.1:5000/page?name=../../etc/passwd。",
        "回车访问页面。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【Linux 复现结果】", bold=True, size=11)
    add_para(doc,
             "成功读取系统用户文件 /etc/passwd，展示服务器所有系统账号、用户权限信息，"
             "如 root、www-data、mysql 等系统用户列表。",
             first_indent=0.74)

    add_para(doc, "【Windows 系统复现步骤】", bold=True, size=11)
    steps = [
        "首页地址栏输入：http://127.0.0.1:5000/page?name=../../Windows/System32/drivers/etc/hosts。",
        "访问页面查看内容。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"步骤 {i}：{step}", first_indent=0.74)

    add_para(doc, "【Windows 复现结果】", bold=True, size=11)
    add_para(doc,
             "读取本地 hosts 文件，获取内网域名、IP 映射信息，服务器本地敏感文件泄露。",
             first_indent=0.74)

    add_para(doc, "攻击原理示意图：", bold=True, size=11)
    add_code_block(doc,
        "Linux Payload:\n"
        "name=../../etc/passwd\n"
        "拼接路径: pages/../../etc/passwd\n"
        "规范化后: /etc/passwd  ← 完全跳出项目目录\n\n"
        "Windows Payload:\n"
        "name=../../Windows/System32/drivers/etc/hosts\n"
        "拼接路径: pages/../../Windows/System32/drivers/etc/hosts\n"
        "规范化后: C:/Windows/System32/drivers/etc/hosts"
    )

    add_heading_styled(doc, "6.2  修复方案", level=2)

    add_para(doc, "【修复核心】", bold=True, size=11)
    add_para(doc,
             "四层防御共同作用：黑名单拦截 ../ 和 / -> 白名单拦截特殊字符 -> "
             "规范化路径 -> 目录归属校验（强制在 pages 内）。",
             first_indent=0.74)

    add_para(doc, "【修复原理】", bold=True, size=11)
    add_para(doc,
             "/etc/passwd 包含 / 和 ..，第 1 层黑名单直接拦截；即使通过编码绕过，"
             "第 3 层规范化后的路径 /etc/passwd 不以项目 pages 目录开头，第 4 层目录归属校验直接拒绝。"
             "多层防御互补，单一绕过无法突破所有防线。",
             first_indent=0.74)

    add_para(doc, "防护链路：", bold=True, size=11)
    add_code_block(doc,
        "name=../../etc/passwd\n"
        "-> 第 1 层: 检测到 '..' 和 '/' → 拦截\n\n"
        "name=..\\..\\Windows\\hosts\n"
        "-> 第 1 层: 检测到 '..' 和 '\\' → 拦截\n\n"
        "即使全部绕过（极端情况）:\n"
        "-> 第 3 层: 规范化路径 = C:/Windows/hosts\n"
        "-> 第 4 层: 不以 pages 目录开头 → 拦截\n"
        "-> 第 4 层: 不以 .html 结尾 → 拦截"
    )

    # ═══════════════════════════════════════════
    # 七、修复汇总
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "七、修复方案汇总与安全对比", level=1)
    add_separator(doc)

    make_table(doc,
               ["漏洞", "攻击示例", "修复前", "修复后"],
               [
                   ["读取 app.py 源码", "?name=../app.py", "直接输出源码", "黑名单拦截 .."],
                   ["读取模板文件", "?name=../templates/base", "自动补 .html 后读取", "黑名单 + 目录校验"],
                   ["读取配置文件", "?name=../config", "直接读取无后缀文件", "后缀 .html 强制校验"],
                   ["读取系统文件", "?name=../../etc/passwd", "读取系统用户列表", "黑名单 + 目录校验"],
               ],
               col_widths=[3, 4, 3.5, 3.5])

    make_table(doc,
               ["安全维度", "修复前", "修复后"],
               [
                   ["路径穿越字符过滤", "无过滤", "黑名单拦截 .. / \\ ~"],
                   ["页面名称校验", "无校验", "白名单 [a-zA-Z0-9_\\-.]"],
                   ["路径规范化", "未使用", "os.path.abspath() 强制规范化"],
                   ["目录归属校验", "无限制", "强制在 pages 目录内"],
                   ["文件后缀限制", "无限制", "仅 .html 文件"],
                   ["自动补 .html 逻辑", "可被利用扩大攻击面", "执行后仍受目录和后缀校验"],
                   ["错误信息", "无提示 / 直接报错", "统一返回「页面不存在」"],
               ],
               col_widths=[4, 4, 6])

    # ═══════════════════════════════════════════
    # 八、修复后验证
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "八、修复后安全验证", level=1)
    add_separator(doc)

    make_table(doc,
               ["测试用例", "预期结果", "实际结果", "结论"],
               [
                   ["正常加载 help 页面", "显示帮助中心内容", "显示帮助中心", "通过"],
                   ["正常加载 about 页面", "显示关于我们内容", "显示关于我们", "通过"],
                   ["路径穿越 ../app.py", "拦截，返回\"页面不存在\"", "页面不存在", "通过"],
                   ["路径穿越 ../templates/base", "拦截，返回\"页面不存在\"", "页面不存在", "通过"],
                   ["读取系统文件 /etc/passwd", "拦截，返回\"页面不存在\"", "页面不存在", "通过"],
                   ["读取配置文件 config", "拦截，返回\"页面不存在\"", "页面不存在", "通过"],
                   ["非法页面名 <script>", "拦截，返回\"页面不存在\"", "页面不存在", "通过"],
                   ["不存在的页面 notexist", "返回\"页面不存在\"", "页面不存在", "通过"],
                   ["无参数请求 /page", "返回\"请指定页面名称\"", "请指定页面名称", "通过"],
               ],
               col_widths=[4, 3.5, 3.5, 1.5])

    add_para(doc, "验证结论：", bold=True, size=11, space_before=8)
    add_para(doc,
             "全部 4 项文件包含/路径遍历漏洞已成功修复。正常业务功能完全保留，"
             "所有攻击向量全部被拦截。修复后系统满足以下安全要求：\n"
             "（1）任何路径穿越字符均被黑名单拦截，无法传入路径拼接逻辑；\n"
             "（2）页面名称白名单严格限制合法字符范围；\n"
             "（3）所有路径经规范化后强制校验目录归属，无法跳出 pages 目录；\n"
             "（4）非 .html 文件一律拒绝读取，阻断敏感文件泄露；\n"
             "（5）错误信息统一化，不泄露服务器路径信息。",
             first_indent=0.74)

    # ═══════════════════════════════════════════
    # 九、总结
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "九、实验总结与心得", level=1)
    add_separator(doc)

    add_para(doc,
             "通过本次动态页面文件包含漏洞复现与修复实验，系统掌握了路径遍历攻击的多种利用手法"
             "及其对应的分层防御策略。",
             first_indent=0.74)

    add_para(doc, "核心认知：", bold=True, size=11)
    add_para(doc,
             "（1）路径遍历攻击的本质是「用户输入控制了文件路径」。只要存在由用户输入拼接文件路径"
             "的逻辑且缺乏有效校验，就存在被利用的可能。防御的核心在于从不信任用户输入，"
             "并在多个层面实施校验。",
             first_indent=0.74)

    add_para(doc,
             "（2）单层防御不足以保证安全。黑名单可以被编码绕过，白名单可能被遗漏字符突破，"
             "目录校验可能被符号链接绕过。本次修复采用的四层防御架构（黑名单 + 白名单 + "
             "规范化 + 目录归属）互为补充，任何单层被突破时仍有其他层提供保护。",
             first_indent=0.74)

    add_para(doc,
             "（3）自动补全逻辑可能扩大攻击面。自动添加 .html 后缀的本意是提升用户体验，"
             "但被攻击者利用后可读取更多类型的文件。安全逻辑必须覆盖所有代码执行路径，"
             "包括异常处理和自动补全分支。",
             first_indent=0.74)

    add_para(doc,
             "（4）路径规范化（os.path.abspath、os.path.realpath）是防御路径穿越的关键技术。"
             "不规范化的路径比较容易被 /.// 等技巧绕过，规范化后 ../ 被实际解析，"
             "可准确判断文件真实位置是否在允许范围内。",
             first_indent=0.74)

    add_para(doc,
             "通过本次实验深刻认识到，文件操作相关功能必须实施严格的路径边界控制。"
             "任何允许用户直接或间接影响文件路径的功能都应被视为高风险功能，"
             "必须在设计阶段就纳入安全校验。",
             first_indent=0.74)

    # 结尾
    doc.add_paragraph()
    add_separator(doc)

    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run("— 报告完 —")
    set_run_font(run, size=11, color=(150, 150, 150))

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "本报告仅供安全教学与技术交流使用。报告中涉及的漏洞代码已全部修复，\n"
        "请勿将存在漏洞的代码版本部署到任何生产环境。"
    )
    set_run_font(run, size=9, color=(180, 180, 180))

    doc.save(output_path)
    print(f"报告已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "动态页面文件包含漏洞_路径遍历复现及修复报告.docx")
    generate_report(out)
